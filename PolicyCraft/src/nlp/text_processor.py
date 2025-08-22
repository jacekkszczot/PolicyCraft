"""
Text Processing Module for PolicyCraft AI Policy Analysis Platform.
This module provides lightweight text processing capabilities for policy documents,
including text extraction from various file formats (PDF, DOCX, TXT) and basic
natural language processing tasks. The implementation uses regex-based tokenization
to minimize external dependencies while maintaining good performance.

Key Features:
- Support for multiple document formats (PDF, DOCX, TXT)
- Regex-based tokenization without NLTK dependency
- Text cleaning and normalization
- Basic text statistics and analysis
- Policy-specific term handling

Author: Jacek Robert Kszczot
Project: MSc Data Science & AI - COM7016
University: Leeds Trinity University
"""

import os
import re
import logging
import sys  # might need this for debugging later
from pathlib import Path
from typing import Optional, Dict, List

# Error messages and status constants
ERROR_FILE_NOT_FOUND = "File not found: {file_path}"
ERROR_UNSUPPORTED_FORMAT = "Unsupported file format: {file_extension}"
ERROR_EXTRACTION = "Could not extract text from {file_path}: {error}"
ERROR_PDF_EXTRACTION = "{method} extraction failed: {error}"
ERROR_CONTRACTION = "Could not expand contractions: {error}"

# Status messages for logging
STATUS_PDF_EXTRACTED = "PDF text extracted using {method}: {char_count} characters"
STATUS_DOCX_EXTRACTED = "DOCX text extracted: {char_count} characters"
STATUS_TXT_EXTRACTED = "TXT file read with {encoding}: {char_count} characters"
STATUS_TEXT_CLEANED = "Text cleaned: {char_count} characters"
STATUS_TOKENISED_SENTENCES = "Tokenised into {count} sentences"  # British spelling
STATUS_TOKENISED_WORDS = "Tokenised into {count} words"          # British spelling
STATUS_STATS_CALCULATED = "Text statistics calculated: {stats}"

# Debug messages
DEBUG_PDF_EXTRACTED = "PyPDF2 extracted {char_count} characters"
DEBUG_PDFPLUMBER_EXTRACTED = "pdfplumber extracted {char_count} characters"

# Output formatting
OUTPUT_ORIGINAL_TEXT = "Original text ({char_count} chars):"
OUTPUT_CLEANED_TEXT = "\nCleaned text ({char_count} chars):"
OUTPUT_STAT_ITEM = "  {key}: {value}"
OUTPUT_SENTENCE_ITEM = "  {index}. {sentence}"
OUTPUT_WORDS_PREVIEW = "\nFirst 10 words: {words}"
OUTPUT_TEXT_PREVIEW = "\nPreview (100 chars): {preview}"

# PDF processing libraries - had some issues with these initially
try:
    import pypdf
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PDF libraries not available. Install with: pip install PyPDF2 pdfplumber")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: DOCX library not available. Install with: pip install python-docx")

try:
    import contractions
    CONTRACTIONS_AVAILABLE = True
except ImportError:
    CONTRACTIONS_AVAILABLE = False
    print("Warning: contractions library not available. Install with: pip install contractions")

logger = logging.getLogger(__name__)

class TextProcessor:
    """
    Simplified text processing for policy documents.
    Uses regex-based tokenization to avoid dependency issues.
    """
    
    def __init__(self):
        """Initialise text processor with configuration."""
        self.supported_formats = {'.pdf', '.docx', '.txt', '.doc'}
        
        # Common English stopwords (simplified list)
        self.stop_words = {
            'i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', 'your', 'yours',
            'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', 'her', 'hers',
            'herself', 'it', 'its', 'itself', 'they', 'them', 'their', 'theirs', 'themselves',
            'what', 'which', 'who', 'whom', 'this', 'that', 'these', 'those', 'am', 'is', 'are',
            'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does',
            'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until',
            'while', 'of', 'at', 'by', 'for', 'with', 'through', 'during', 'before', 'after',
            'above', 'below', 'up', 'down', 'in', 'out', 'on', 'off', 'over', 'under', 'again',
            'further', 'then', 'once'
        }
        
        # Policy-specific terms to preserve
        self.policy_terms = {
            'ai', 'artificial', 'intelligence', 'policy', 'ethics', 'bias',
            'transparency', 'accountability', 'fairness', 'privacy', 'data',
            'algorithm', 'model', 'training', 'deployment', 'governance',
            'compliance', 'regulation', 'guideline', 'framework', 'principle'
        }
        
        # Remove policy terms from stopwords
        self.stop_words = self.stop_words - self.policy_terms
        
        logger.info("TextProcessor initialised successfully")

    def extract_text_from_file(self, file_path: str) -> Optional[str]:
        """Extract text from various file formats with fallback methods."""
        file_path = Path(file_path)
        if not file_path.exists():
            logger.warning(ERROR_FILE_NOT_FOUND.format(file_path=file_path))
            return None
            
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            logger.warning(ERROR_UNSUPPORTED_FORMAT.format(file_extension=file_extension))
            return None
            
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
                
        except Exception as e:
            logger.warning(ERROR_EXTRACTION.format(file_path=file_path, error=str(e)))
            return None
            
        return None

    def _extract_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF using multiple methods including OCR for scanned documents."""
        if not PDF_AVAILABLE:
            logger.warning("PDF libraries not available")
            return None
            
        text_pypdf2 = ""
        text_pdfplumber = ""
        text_ocr = ""
        
        # Method 1: PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page in pdf_reader.pages:
                    text_pypdf2 += page.extract_text() + "\n"
        except Exception as e:
            logger.warning(ERROR_PDF_EXTRACTION.format(method="PyPDF2", error=str(e)))
            
        # Method 2: pdfplumber
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_pdfplumber += page_text + "\n"
        except Exception as e:
            logger.warning(ERROR_PDF_EXTRACTION.format(method="pdfplumber", error=str(e)))
            
        # Method 3: OCR (if standard methods fail or extract minimal text)
        if not text_pypdf2.strip() and not text_pdfplumber.strip() or \
           (len(text_pypdf2.strip()) < 100 and len(text_pdfplumber.strip()) < 100):
            try:
                text_ocr = self._extract_text_with_ocr(file_path)
                if text_ocr:
                    logger.info(f"OCR extraction successful: {len(text_ocr)} characters")
            except Exception as e:
                logger.warning(f"OCR extraction failed: {e}")
            
        # Choose the best result
        texts = [
            (text_pypdf2, "PyPDF2"),
            (text_pdfplumber, "pdfplumber"), 
            (text_ocr, "OCR")
        ]
        
        best_text, method = max(texts, key=lambda x: len(x[0].strip()))
            
        if best_text.strip():
            logger.info(STATUS_PDF_EXTRACTED.format(
                method=method,
                char_count=len(best_text)
            ))
            return best_text
        else:
            logger.warning("No text could be extracted from PDF using any method")
            return None

    def _extract_from_docx(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX files."""
        if not DOCX_AVAILABLE:
            logger.warning("DOCX library not available")
            return None
            
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if text.strip():
                logger.info(STATUS_DOCX_EXTRACTED.format(char_count=len(text)))
                return text
            else:
                logger.warning("No text could be extracted from DOCX file")
                return None
            
        except Exception as e:
            logger.warning(ERROR_EXTRACTION.format(
                file_path=file_path,
                error=str(e)
            ))
            return None

    def _extract_from_txt(self, file_path: str) -> Optional[str]:
        """Extract text from TXT files with encoding detection."""
        encodings = ['utf-8', 'latin-1', 'windows-1252', 'ascii']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    if text.strip():
                        logger.info(STATUS_TXT_EXTRACTED.format(
                            encoding=encoding,
                            char_count=len(text)
                        ))
                        return text
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(ERROR_EXTRACTION.format(
                    file_path=file_path,
                    error=str(e)
                ))
                return None
                
        logger.warning("Could not read file with any encoding: %s", file_path)
        return None

    def clean_text(self, text: str) -> str:
        """Comprehensive text cleaning and preprocessing."""
        if not text:
            return ""
            
        
        # Compile regex patterns for better performance
        regex_patterns = {
            'whitespace': re.compile(r'\s+'),
            'newline': re.compile(r'\n+'),
            'special_chars': re.compile(r'[^\w\s\.\,\!\?\;\:\-\(\)]'),
            'punctuation_before': re.compile(r'\s+([\.!?])'),
            'punctuation_after': re.compile(r'([\.!?])\s*'),
            'url': re.compile(
                r'https?://'  # http:// or https://
                r'(?:(?:[A-Z0-9](?:[A-Z0-9-]{0,61}[A-Z0-9])?\.)+[A-Z]{2,}\.?|'  # domain
                r'localhost|'  # localhost
                r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})'  # IP
                r'(?::\d+)?'  # optional port
                r'(?:(?:/|\?)[^\s]*)?',  # path and query params
                re.IGNORECASE
            ),
            'email': re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'),
            'multi_dots': re.compile(r'\.{2,}'),
            'multi_question': re.compile(r'\?{2,}'),
            'multi_exclaim': re.compile(r'!{2,}'),
            'multi_hyphen': re.compile(r'-{2,}'),
            'trailing_whitespace': re.compile(r'\s+$'),
            'leading_whitespace': re.compile(r'^\s+'),
            'multi_space': re.compile(r' {2,}'),
        }
        
        try:
            # Basic cleaning - normalize whitespace and newlines
            text = regex_patterns['whitespace'].sub(' ', text)
            text = regex_patterns['newline'].sub('\n', text)
            
            # Expand contractions if available
            if CONTRACTIONS_AVAILABLE:
                try:
                    text = contractions.fix(text)
                except Exception as e:
                    logger.warning(f"Failed to expand contractions: {e}")
            
            # Remove special characters but preserve sentence structure
            text = regex_patterns['special_chars'].sub(' ', text)
            
            # Fix spacing around punctuation
            text = regex_patterns['punctuation_before'].sub(r'\1', text)  # Remove spaces before punctuation
            text = regex_patterns['punctuation_after'].sub(r'\1 ', text)  # Ensure single space after punctuation
            
            # Remove URLs and emails
            text = regex_patterns['url'].sub('', text)
            text = regex_patterns['email'].sub('', text)
            
            # Normalize excessive punctuation
            text = regex_patterns['multi_dots'].sub('.', text)  # Replace multiple dots with single dot
            text = regex_patterns['multi_question'].sub('?', text)  # Replace multiple question marks with one
            text = regex_patterns['multi_exclaim'].sub('!', text)  # Replace multiple exclamations with one
            text = regex_patterns['multi_hyphen'].sub('-', text)  # Replace multiple hyphens with one
            
            # Final cleanup
            text = regex_patterns['trailing_whitespace'].sub('', text)
            text = regex_patterns['leading_whitespace'].sub('', text)
            text = regex_patterns['multi_space'].sub(' ', text)
            
            return text.strip()
            
        except Exception as e:
            logger.error("Error in text preprocessing: %s", str(e))
            return text  # Return partially cleaned text if an error occurs
        
        # Final whitespace normalization
        text = whitespace_pattern.sub(' ', text).strip()
        text = text.strip()
        logger.debug(f"Text cleaned: {len(text)} characters")
        return text

    def tokenize_sentences(self, text: str) -> List[str]:
        """Split text into sentences using regex."""
        if not text:
            return []
            
        # Simple sentence tokenization using regex
        sentences = re.split(r'[.!?]+', text)
        
        # Clean and filter sentences
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 10:  # Filter out very short sentences
                cleaned_sentences.append(sentence)
        
        logger.debug(f"Tokenized into {len(cleaned_sentences)} sentences")
        return cleaned_sentences

    def tokenize_words(self, text: str, remove_stopwords: bool = True) -> List[str]:
        """Split text into words using regex."""
        if not text:
            return []
            
        # Extract words using regex
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        
        # Filter stopwords if requested
        if remove_stopwords:
            words = [word for word in words if word not in self.stop_words]
                
        logger.debug(f"Tokenized into {len(words)} words")
        return words

    def get_text_statistics(self, text: str) -> Dict:
        """Calculate various text statistics."""
        if not text:
            return {
                'character_count': 0,
                'word_count': 0,
                'sentence_count': 0,
                'paragraph_count': 0,
                'avg_sentence_length': 0,
                'avg_word_length': 0
            }
            
        sentences = self.tokenize_sentences(text)
        words = self.tokenize_words(text, remove_stopwords=False)
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        stats = {
            'character_count': len(text),
            'word_count': len(words),
            'sentence_count': len(sentences),
            'paragraph_count': len(paragraphs),
            'avg_sentence_length': round(len(words) / len(sentences), 2) if sentences else 0,
            'avg_word_length': round(sum(len(word) for word in words) / len(words), 2) if words else 0
        }
        
        logger.debug(f"Text statistics calculated: {stats}")
        return stats

    def preview_text(self, text: str, max_length: int = 500) -> str:
        """Generate a preview of the text."""
        if not text:
            return ""
            
        if len(text) <= max_length:
            return text
            
        # Find a good breaking point
        preview = text[:max_length]
        last_period = preview.rfind('.')
        
        if last_period > max_length * 0.7:
            preview = preview[:last_period + 1]
        
        return preview + "..." if len(text) > len(preview) else preview

    def _extract_text_with_ocr(self, file_path: str) -> Optional[str]:
        """Extract text from PDF using OCR for scanned documents."""
        try:
            import pytesseract
            from pdf2image import convert_from_path
            from PIL import Image
            
            # Convert PDF pages to images
            images = convert_from_path(file_path, dpi=300)
            
            if not images:
                logger.warning("No images could be extracted from PDF for OCR")
                return None
            
            # Extract text from each image using OCR
            ocr_text = ""
            for i, image in enumerate(images):
                try:
                    # Convert PIL image to format suitable for tesseract
                    page_text = pytesseract.image_to_string(image, lang='eng')
                    if page_text.strip():
                        ocr_text += f"--- Page {i+1} ---\n{page_text.strip()}\n\n"
                        logger.info(f"OCR extracted {len(page_text.strip())} characters from page {i+1}")
                except Exception as e:
                    logger.warning(f"OCR failed for page {i+1}: {e}")
                    continue
            
            if ocr_text.strip():
                logger.info(f"OCR extraction completed: {len(ocr_text)} total characters")
                return ocr_text
            else:
                logger.warning("OCR extraction produced no text")
                return None
                
        except ImportError as e:
            logger.warning(f"OCR libraries not available: {e}")
            return None
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return None


# Test the processor
if __name__ == "__main__":
    processor = TextProcessor()
    print("TextProcessor test completed successfully!")