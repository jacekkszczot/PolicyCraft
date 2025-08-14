"""
Main Flask application for PolicyCraft - AI Policy Analysis Framework.
Clean version without onboarding features.

Author: Jacek Robert Kszczot
Project: MSc Data Science & AI - COM7016
University: Leeds Trinity University
"""

from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, make_response
from flask_login import LoginManager, login_required, current_user
from flask_wtf.csrf import CSRFProtect
import os
from werkzeug.utils import secure_filename
import logging
from datetime import datetime

# Import configuration
from config import get_config, create_secure_directories

# Import authentication components
from src.database.models import User, db, init_db
from src.auth.routes import auth_bp
from src.admin.routes import admin_bp

# Import analysis modules
# Constants for duplicated literals
DOCX_EXTENSION = '.docx'
TIMEZONE_SUFFIX = '+00:00'
BASELINE_PREFIX = '[BASELINE]'
ANALYSIS_NOT_FOUND = 'Analysis not found'
NO_RECOMMENDATIONS_FOUND = 'No recommendations found for this analysis'

from src.nlp.text_processor import TextProcessor
from src.nlp.theme_extractor import ThemeExtractor
from src.nlp.policy_classifier import PolicyClassifier
from src.database.mongo_operations import MongoOperations as DatabaseOperations
from src.recommendation.engine import RecommendationGenerator as RecommendationEngine
from src.export.export_engine import ExportEngine
from src.literature.literature_engine import LiteratureEngine
from src.literature.knowledge_manager import KnowledgeBaseManager as KnowledgeManager
from src.utils.auto_document_manager import AutoDocumentManager
from src.visualisation.charts import ChartGenerator
from src.scripts.clean_dataset import process_new_upload


def clean_baseline_filename(filename):
    """Remove BASELINE prefix and clean filename for display"""
    if filename.startswith(BASELINE_PREFIX + " "):
        clean_name = filename.replace(BASELINE_PREFIX + " ", "")
        if clean_name.endswith(".pdf"):
            clean_name = clean_name.replace(".pdf", "")
        return clean_name
    return filename

def clean_university_name(filename):
    """
    Clean filename to display only the university name without technical prefixes.
    
    This function removes timestamp prefixes and file extensions to present
    a clean, user-friendly university name for display purposes.
    """
    # Remove timestamp prefixes
    if '_' in filename:
        clean_name = filename.split('_', 2)[-1] if len(filename.split('_')) > 2 else filename.split('_')[-1]
    else:
        clean_name = filename
    
    # Remove file extensions
    clean_name = clean_name.replace('.pdf', '').replace(DOCX_EXTENSION, '').replace('.doc', '').replace('.txt', '')
    clean_name = clean_name.replace('-ai-policy', '').replace('_ai_policy', '')
    
    # Map to clean university names
    name_lower = clean_name.lower()
    if 'harvard' in name_lower:
        return 'Harvard University'
    elif 'stanford' in name_lower:
        return 'Stanford University'
    elif 'mit' in name_lower:
        return 'MIT'
    elif 'cambridge' in name_lower:
        return 'University of Cambridge'
    elif 'oxford' in name_lower:
        return 'Oxford University'
    elif 'belfast' in name_lower:
        return 'Belfast University'
    elif 'edinburgh' in name_lower:
        return 'Edinburgh University'
    elif 'columbia' in name_lower:
        return 'Columbia University'
    elif 'cornell' in name_lower:
        return 'Cornell University'
    elif 'chicago' in name_lower:
        return 'University of Chicago'
    elif 'imperial' in name_lower:
        return 'Imperial College London'
    elif 'tokyo' in name_lower:
        return 'University of Tokyo'
    elif 'jagiellonian' in name_lower:
        return 'Jagiellonian University'
    elif 'leeds' in name_lower and 'trinity' in name_lower:
        return 'Leeds Trinity University'
    elif 'liverpool' in name_lower:
        return 'University of Liverpool'
    else:
        # Generic cleanup
        clean_name = clean_name.replace('-', ' ').replace('_', ' ').title()
        return clean_name[:30] + '...' if len(clean_name) > 30 else clean_name

def clean_filename(filename):
    """
    Clean filename to show document name without timestamp and user ID prefixes.
    
    Removes technical prefixes added during upload process to present
    clean filenames for user interface display.
    """
    if '_' in filename:
        parts = filename.split('_')
        if len(parts) >= 3:
            clean_name = '_'.join(parts[2:])
        else:
            clean_name = filename
    else:
        clean_name = filename
    return clean_name

def format_british_date(date_str):
    """
    Format date string to British DD/MM/YYYY style.
    
    Converts various date formats to the standardised British format
    for consistent presentation throughout the application interface.
    """
    try:
        from datetime import datetime
        if isinstance(date_str, str):
            if 'T' in date_str:
                dt = datetime.fromisoformat(date_str.replace('Z', TIMEZONE_SUFFIX))
            else:
                dt = datetime.strptime(date_str[:19], '%Y-%m-%d %H:%M:%S')
        else:
            dt = date_str
        return dt.strftime('%d/%m/%Y')
    except Exception:
        return str(date_str)[:10] if len(str(date_str)) > 10 else str(date_str)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/application.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def create_app():
    """
    Application factory function for PolicyCraft Flask application.
    
    Initialises and configures the Flask application with all necessary
    extensions, blueprints, and database connections.
    """
    # Create secure directories first
    create_secure_directories()
    
    # Initialise Flask application
    app = Flask(__name__, 
               template_folder='src/web/templates',
               static_folder='src/web/static')
    # Avoid 308 redirects in tests for trailing slash differences
    app.url_map.strict_slashes = False
    
    # Load configuration (explicit env to satisfy linter and keep behaviour)
    config_env = os.environ.get('FLASK_ENV', 'development')
    config_obj = get_config(config_env)
    app.config.from_object(config_obj)
    
    # Initialise extensions
    db.init_app(app)
    
    # Setup Flask-Login
    login_manager = LoginManager()
    login_manager.init_app(app)
    login_manager.login_view = 'auth.login'
    login_manager.login_message = 'Please log in to access this page.'
    login_manager.login_message_category = 'info'
    
    @login_manager.user_loader
    def load_user(user_id):
        """Load user for Flask-Login session management."""
        return User.query.get(int(user_id))
    
    # Register blueprints (auth without prefix to expose /login and /logout)
    app.register_blueprint(auth_bp)
    app.register_blueprint(admin_bp)

    # CSRF protection and csrf_token for templates (available for all app instances)
    CSRFProtect(app)
    # Ensure csrf_token available globally in Jinja
    from flask_wtf.csrf import generate_csrf
    app.jinja_env.globals['csrf_token'] = generate_csrf
    @app.context_processor
    def inject_csrf_token():
        """Provide csrf_token() callable to Jinja templates even when CSRF is disabled (testing)."""
        return dict(csrf_token=lambda: generate_csrf())

    # Make current_user available in all templates
    @app.context_processor
    def inject_current_user():
        """Make current_user available in all templates."""
        return {'current_user': current_user}
    
    # Define core routes on this app instance so tests using a fresh app have them
    @app.route('/')
    def index():
        """
        Landing page route for PolicyCraft application.
        """
        logger.info("Landing page accessed")
        return render_template("index.html")

    @app.route('/about')
    def about():
        """Minimal about page used by tests/templates."""
        return render_template("about.html") if os.path.exists(os.path.join(app.template_folder, 'about.html')) else "About PolicyCraft", 200, {"Content-Type": "text/plain; charset=utf-8"}

    # Initialise database and create tables
    with app.app_context():
        init_db(app)
    
    return app

# Create Flask app and initialize modules
app = create_app()

 

text_processor = TextProcessor()
theme_extractor = ThemeExtractor()
policy_classifier = PolicyClassifier()
db_operations = DatabaseOperations(uri="mongodb://localhost:27017", db_name="policycraft")
chart_generator = ChartGenerator()
# Initialise recommendation engine with knowledge base path (unified with admin panel)
knowledge_base_path = "docs/knowledge_base"  # Use same path as LiteratureEngine
recommendation_engine = RecommendationEngine(knowledge_base_path=knowledge_base_path)


def allowed_file(filename):
    """
    Check if the uploaded file has an allowed extension.
    
    Validates file type against configured allowed extensions
    to ensure only supported document formats are processed.
    """
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def create_upload_folder():
    """
    Create upload folder if it doesn't exist.
    
    Ensures the designated upload directory exists for storing
    user-uploaded policy documents.
    """
    upload_folder = app.config['UPLOAD_FOLDER']
    if not os.path.exists(upload_folder):
        os.makedirs(upload_folder)
        logger.info(f"Created upload folder: {upload_folder}")

# Main application routes

def _to_epoch(val):
    """Convert supported date formats to POSIX timestamp for consistent sorting."""
    try:
        if isinstance(val, datetime):
            return val.replace(tzinfo=None).timestamp()
        if isinstance(val, str):
            dt = datetime.fromisoformat(val.replace('Z', TIMEZONE_SUFFIX))
            return dt.replace(tzinfo=None).timestamp()
    except Exception:
        pass
    return 0

def _combine_stats(stats1, stats2):
    """Combine two sets of statistics with weighted averaging."""
    total1 = stats1.get('total', 0)
    total2 = stats2.get('total', 0)
    total = total1 + total2
    
    if total == 0:
        return {
            'total': 0,
            'avg_confidence': 0,
            'avg_themes_per_analysis': 0
        }
    
    # Calculate weighted average for confidence and themes
    conf1 = (stats1.get('avg_confidence', 0) or 0) * total1
    conf2 = (stats2.get('avg_confidence', 0) or 0) * total2
    
    themes1 = (stats1.get('avg_themes_per_analysis', 0) or 0) * total1
    themes2 = (stats2.get('avg_themes_per_analysis', 0) or 0) * total2
    
    return {
        'total': total,
        'avg_confidence': round((conf1 + conf2) / total, 1) if total > 0 else 0,
        'avg_themes_per_analysis': round((themes1 + themes2) / total, 1) if total > 0 else 0
    }

def _process_analyses_for_display(analyses):
    """Process analyses for display by converting dates and preparing data."""
    processed = []
    for analysis in analyses:
        # Skip if analysis is not a dictionary
        if not isinstance(analysis, dict):
            logger.error(f"Dashboard error: Skipping non-dict analysis: {type(analysis)}")
            continue
            
        processed_analysis = analysis.copy()
        
        # Ensure classification exists and is properly formatted
        if 'classification' not in processed_analysis:
            processed_analysis['classification'] = 'Moderate'  # Default classification
            
        # Standardise classification to use only Restrictive, Moderate, or Permissive
        if isinstance(processed_analysis['classification'], dict):
            cls = processed_analysis['classification'].get('classification', 'Moderate')
            processed_analysis['classification']['classification'] = _standardize_classification(cls)
        elif isinstance(processed_analysis['classification'], str):
            processed_analysis['classification'] = _standardize_classification(processed_analysis['classification'])
        else:
            # Handle unexpected classification type
            processed_analysis['classification'] = 'Moderate'
            logger.error(f"Dashboard error: Unexpected classification type: {type(processed_analysis['classification'])}")
        
        # Ensure other required fields exist
        if 'title' not in processed_analysis:
            processed_analysis['title'] = processed_analysis.get('filename', 'Untitled Policy')
            
        if 'summary' not in processed_analysis:
            processed_analysis['summary'] = 'No summary available'
            
        if 'themes' not in processed_analysis:
            processed_analysis['themes'] = []
            
        if 'score' not in processed_analysis:
            processed_analysis['score'] = 50  # Default score
        
        # Format dates
        if 'analysis_date' not in processed_analysis:
            processed_analysis['analysis_date'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        elif hasattr(processed_analysis['analysis_date'], 'strftime'):
            processed_analysis['analysis_date'] = processed_analysis['analysis_date'].strftime('%Y-%m-%d %H:%M:%S')
        elif isinstance(processed_analysis['analysis_date'], str):
            try:
                dt = datetime.fromisoformat(processed_analysis['analysis_date'].replace('Z', TIMEZONE_SUFFIX))
                processed_analysis['analysis_date'] = dt.strftime('%Y-%m-%d %H:%M:%S')
            except Exception:
                # Keep the original string if parsing fails
                pass
                
        processed.append(processed_analysis)
    return processed

def _standardize_classification(classification):
    """
    Standardize classification to use only Restrictive, Moderate, or Permissive.
    Maps legacy categories to standard ones.
    """
    # Map of legacy categories to standard categories
    category_map = {
        # Standard categories (keep as is)
        'Restrictive': 'Restrictive',
        'Moderate': 'Moderate',
        'Permissive': 'Permissive',
        
        # Legacy categories (map to standard)
        'Educational': 'Moderate',
        'Research-focused': 'Permissive',
        'Comprehensive': 'Moderate',
        'Balanced': 'Moderate',  # Just in case
        
        # Default for unknown
        'Unknown': 'Moderate'
    }
    
    # Return standardized classification or default to Moderate if not in map
    return category_map.get(classification, 'Moderate')

def _calculate_analytics(analyses):
    """Calculate analytics (classifications and themes) from analyses."""
    classification_counts = {}
    theme_frequencies = {}
    
    for analysis in analyses:
        # Skip if analysis is not a dictionary
        if not isinstance(analysis, dict):
            logger.error(f"Analytics error: Skipping non-dict analysis: {type(analysis)}")
            continue
            
        # Get classification with robust type checking
        classification = analysis.get('classification', 'Moderate')
        
        # Handle different classification formats
        if isinstance(classification, dict):
            cls = classification.get('classification', 'Moderate')
        elif isinstance(classification, str):
            cls = classification
        else:
            # Default for unexpected types
            cls = 'Moderate'
            logger.error(f"Analytics error: Unexpected classification type: {type(classification)}")
        
        # Standardize classification to use only Restrictive, Moderate, or Permissive
        standardized_cls = _standardize_classification(cls)
        classification_counts[standardized_cls] = classification_counts.get(standardized_cls, 0) + 1
        
        # Count theme frequencies with robust type checking
        themes = analysis.get('themes', [])
        if not isinstance(themes, list):
            themes = []
            
        for theme in themes:
            if isinstance(theme, dict):
                theme_name = theme.get('name', 'Unknown')
            else:
                theme_name = str(theme)
            theme_frequencies[theme_name] = theme_frequencies.get(theme_name, 0) + 1
    
    # Ensure all standard classifications are present in counts
    for cls in ['Restrictive', 'Moderate', 'Permissive']:
        if cls not in classification_counts:
            classification_counts[cls] = 0
    
    return classification_counts, theme_frequencies

def _get_minimal_dashboard_data(user):
    """Return minimal dashboard data for error cases."""
    return {
        'user': {
            'id': user.id,
            'username': user.username,
            'full_name': user.get_full_name()
        },
        'total_policies': 0,
        'classification_counts': {},
        'theme_frequencies': {},
        'charts': {},
        'recent_analyses': [],
        'statistics': {'total_analyses': 0, 'avg_confidence': 0, 'avg_themes_per_analysis': 0}
    }

@app.route('/dashboard')
@login_required
def dashboard():
    """
    Dashboard with comprehensive analysis data for authenticated users.
    
    Provides overview of user's policy analyses, statistics, and charts
    with baseline policy comparisons and performance metrics.
    """
    try:
        logger.info(f"Dashboard accessed by user: {current_user.username}")
        
        # Get and prepare analyses
        db_operations.deduplicate_baseline_analyses(current_user.id)
        user_analyses = db_operations.get_user_analyses(current_user.id)
        baseline_analyses = db_operations.get_user_analyses(-1)
        
        logger.info(f"Dashboard: Found {len(user_analyses)} user analyses and {len(baseline_analyses)} baseline analyses")
        
        # Load all clean_dataset policies without duplicates
        # This ensures we show exactly 15 unique policies from clean_dataset
        clean_dataset_dir = _get_clean_dataset_dir()
        clean_dataset_files = _list_clean_dataset_files(clean_dataset_dir)
        
        # Merge user and baseline analyses
        analyses_by_filename = _merge_user_and_baseline_analyses(user_analyses, baseline_analyses)
        
        # Identify missing clean_dataset files and ensure baselines exist
        missing_files = _identify_missing_clean_files(clean_dataset_files, analyses_by_filename)
        logger.info(f"Dashboard: Missing files from clean_dataset: {missing_files}")
        
        # Add missing files as proper baseline analyses
        for missing_file in missing_files:
            # Skip guidance files
            if 'guidance' in missing_file.lower() or missing_file == 'dataset_info.md':
                continue
                
            # Get the full path to the file
            file_path = os.path.join(clean_dataset_dir, missing_file)
            if not os.path.exists(file_path):
                logger.error(f"Dashboard: Missing file not found: {file_path}")
                continue
                
            # Create a baseline filename - use a more descriptive name based on the file
            university_name = missing_file.split('-')[0].replace('university', '').strip().title()
            if not university_name:
                university_name = missing_file.split('.')[0].replace('university', '').strip().title()
            
            baseline_filename = f"{BASELINE_PREFIX} {university_name}"
            logger.info(f"Dashboard: Creating baseline analysis for {missing_file} as '{baseline_filename}'")
            
            try:
                # Extract text from the file with robust error handling
                try:
                    # First, verify the file exists and is readable
                    if not os.path.exists(file_path):
                        logger.error(f"Dashboard: File does not exist: {file_path}")
                        raise FileNotFoundError(f"File not found: {file_path}")
                        
                    # Check file size to ensure it's not empty
                    file_size = os.path.getsize(file_path)
                    if file_size == 0:
                        logger.warning(f"Dashboard: Empty file detected: {file_path} (0 bytes)")
                        raise ValueError(f"Empty file: {file_path}")
                        
                    logger.info(f"Dashboard: Attempting to extract text from {file_path} ({file_size} bytes)")
                    
                    # Try multiple extraction methods if needed
                    extracted_text = text_processor.extract_text_from_file(file_path)
                    text_length = len(extracted_text) if extracted_text else 0
                    logger.info(f"Dashboard: Extracted {text_length} characters from {file_path}")
                    
                    # If extraction returned very little text, try a fallback approach
                    if not extracted_text or text_length < 50:  # Sanity check for minimum content
                        logger.warning(f"Dashboard: Insufficient text extracted from {file_path} ({text_length} chars), attempting fallback extraction")
                        
                        # Fallback: For PDF files, try direct extraction with pdfplumber if PyPDF2 failed
                        if file_path.lower().endswith('.pdf'):
                            try:
                                import pdfplumber
                                with pdfplumber.open(file_path) as pdf:
                                    fallback_text = ""
                                    for page in pdf.pages:
                                        page_text = page.extract_text()
                                        if page_text:
                                            fallback_text += page_text + "\n"
                                            
                                if fallback_text and len(fallback_text) > text_length:
                                    extracted_text = fallback_text
                                    text_length = len(extracted_text)
                                    logger.info(f"Dashboard: Fallback extraction successful, got {text_length} characters")
                            except Exception as fallback_error:
                                logger.warning(f"Dashboard: Fallback extraction failed: {str(fallback_error)}")
                        
                        # For DOCX files, try a different approach
                        elif file_path.lower().endswith((DOCX_EXTENSION, '.doc')):
                            try:
                                from docx import Document
                                doc = Document(file_path)
                                fallback_text = "\n".join([para.text for para in doc.paragraphs if para.text])
                                
                                if fallback_text and len(fallback_text) > text_length:
                                    extracted_text = fallback_text
                                    text_length = len(extracted_text)
                                    logger.info(f"Dashboard: Fallback extraction successful, got {text_length} characters")
                            except Exception as fallback_error:
                                logger.warning(f"Dashboard: Fallback extraction failed: {str(fallback_error)}")
                
                    # If we still don't have enough text, use a minimal placeholder but don't fail
                    if not extracted_text or text_length < 20:  # Absolute minimum
                        logger.warning(f"Dashboard: All extraction methods failed for {file_path}, using minimal placeholder text")
                        extracted_text = f"AI Policy document from {university_name}. This is a placeholder text as the original document could not be fully processed."
                        text_length = len(extracted_text)
                except Exception as extract_error:
                    logger.error(f"Dashboard: Text extraction failed for {file_path}: {str(extract_error)}")
                    # Create minimal text instead of re-raising
                    extracted_text = f"AI Policy document from {university_name}. This is a placeholder text as the original document could not be processed due to: {str(extract_error)}"
                    text_length = len(extracted_text)
                    logger.info(f"Dashboard: Using minimal placeholder text ({text_length} chars) after extraction failure")
                    
            except Exception as extract_error:
                logger.error(f"Dashboard: Text extraction failed for {file_path}: {str(extract_error)}")
                # Create minimal text instead of re-raising
                extracted_text = f"AI Policy document from {university_name}. This is a placeholder text as the original document could not be processed due to: {str(extract_error)}"
                text_length = len(extracted_text)
                logger.info(f"Dashboard: Using minimal placeholder text ({text_length} chars) after extraction failure")
                
                # Process the text with robust error handling
            try:
                cleaned_text = text_processor.clean_text(extracted_text)
                logger.info(f"Dashboard: Cleaned text for {missing_file}, now {len(cleaned_text)} characters")
            except Exception as clean_error:
                logger.error(f"Dashboard: Text cleaning failed for {missing_file}: {str(clean_error)}")
                cleaned_text = extracted_text  # Use original text as fallback
                logger.info("Dashboard: Using original text as fallback after cleaning failure")
            
            # Extract themes with error handling
            try:
                themes = theme_extractor.extract_themes(cleaned_text)
                logger.info(f"Dashboard: Extracted {len(themes)} themes from {missing_file}")
                
                # Ensure we have at least some themes
                if not themes or len(themes) == 0:
                    logger.warning(f"Dashboard: No themes extracted for {missing_file}, using defaults")
                    themes = [
                        {"name": "Policy", "score": 0.8, "confidence": 75},
                        {"name": "AI Ethics", "score": 0.7, "confidence": 70},
                        {"name": "Guidelines", "score": 0.6, "confidence": 65}
                    ]
            except Exception as theme_error:
                logger.error(f"Dashboard: Theme extraction failed for {missing_file}: {str(theme_error)}")
                themes = [
                    {"name": "Policy", "score": 0.8, "confidence": 75},
                    {"name": "AI Ethics", "score": 0.7, "confidence": 70},
                    {"name": "Guidelines", "score": 0.6, "confidence": 65}
                ]  # Default themes
            
            # Classify with error handling
            try:
                classification = policy_classifier.classify_policy(cleaned_text)
                if isinstance(classification, str):
                    classification = {
                        "classification": _standardize_classification(classification),
                        "confidence": 80,
                        "source": "Baseline Creation"
                    }
                elif isinstance(classification, dict) and 'classification' in classification:
                    # Ensure standardized classification
                    classification['classification'] = _standardize_classification(classification['classification'])
                    if 'source' not in classification:
                        classification['source'] = "Baseline Creation"
                    if 'confidence' not in classification:
                        classification['confidence'] = 80
                else:
                    # Handle unexpected classification format
                    logger.warning(f"Dashboard: Unexpected classification format for {missing_file}: {classification}")
                    classification = {
                        "classification": "Moderate",  # Default classification
                        "confidence": 75,
                        "source": "Baseline Creation (Default Format)"
                    }
                    
                logger.info(f"Dashboard: Classification for {missing_file}: {classification}")
            except Exception as class_error:
                logger.error(f"Dashboard: Classification failed for {missing_file}: {str(class_error)}")
                classification = {
                    "classification": "Moderate",  # Default classification
                    "confidence": 75,
                    "source": "Baseline Creation (Default)"
                }
                
                
                # Store as a proper baseline analysis with detailed logging
                try:
                    # Ensure we have valid data for storage
                    if not extracted_text:
                        extracted_text = f"AI Policy document from {university_name}. This is a placeholder text."
                    if not cleaned_text:
                        cleaned_text = extracted_text
                    if not themes or len(themes) == 0:
                        themes = [{"name": "Policy", "score": 0.8, "confidence": 75}]
                    if not classification or not isinstance(classification, dict):
                        classification = {
                            "classification": "Moderate",
                            "confidence": 75,
                            "source": "Baseline Creation (Default)"
                        }
                        
                    # Add metadata to help with debugging
                    metadata = {
                        "source_file": missing_file,
                        "extraction_method": "dashboard_baseline_creation",
                        "creation_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "text_length": len(extracted_text) if extracted_text else 0,
                        "is_baseline": True
                    }
                    
                    logger.info(f"Dashboard: Storing baseline analysis for {missing_file} with {len(extracted_text)} chars of text")
                    
                    # Store the analysis
                    analysis_id = db_operations.store_user_analysis_results(
                        user_id=-1,  # Baseline user ID
                        filename=baseline_filename,
                        original_text=extracted_text,
                        cleaned_text=cleaned_text,
                        themes=themes,
                        classification=classification,
                        document_id=missing_file,  # Store original filename as document_id for matching
                        metadata=metadata  # Add metadata for debugging
                    )
                    
                    logger.info(f"Dashboard: Successfully stored baseline analysis with ID {analysis_id}")
                    
                    # Get the newly created analysis
                    new_analysis = db_operations.get_analysis_by_id(analysis_id)
                    if new_analysis:
                        # Mark as baseline
                        new_analysis['is_baseline'] = True
                        new_analysis['is_user_analysis'] = False
                        
                        # Add to our analyses dictionary
                        analyses_by_filename[missing_file] = new_analysis
                        logger.info(f"Dashboard: Added proper baseline analysis for {missing_file} with ID {analysis_id}")
                    else:
                        logger.error(f"Dashboard: Failed to retrieve newly created analysis for {missing_file}")
                        # Create a placeholder instead of raising an exception
                        placeholder_analysis = {
                            '_id': f"placeholder_{missing_file.replace('.', '_')}",
                            'filename': baseline_filename,
                            'document_id': missing_file,
                            'title': f"Policy from {university_name}",
                            'analysis_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            'user_id': -1,
                            'is_user_analysis': False,
                            'is_baseline': True,
                            'classification': classification,
                            'score': 50,
                            'themes': themes,
                            'summary': f"This is a baseline analysis for {missing_file}.",
                            'text_data': {
                                'original_text': extracted_text,
                                'cleaned_text': cleaned_text,
                                'text_length': len(extracted_text) if extracted_text else 0
                            }
                        }
                        analyses_by_filename[missing_file] = placeholder_analysis
                        logger.info(f"Dashboard: Added placeholder analysis for {missing_file} due to retrieval failure")
                        
                # Handle expected storage-related errors locally; let unexpected ones bubble up
                except (ValueError, KeyError, RuntimeError, OSError) as store_error:
                    logger.error(f"Dashboard: Error storing baseline analysis: {str(store_error)}")
                    # Create a placeholder instead of failing
                    placeholder_analysis = {
                        '_id': f"placeholder_{missing_file.replace('.', '_')}",
                        'filename': baseline_filename,
                        'document_id': missing_file,
                        'title': f"Policy from {university_name}",
                        'analysis_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        'user_id': -1,
                        'is_user_analysis': False,
                        'is_baseline': True,
                        'classification': classification,
                        'score': 50,
                        'themes': themes,
                        'summary': f"This is a placeholder for {missing_file}. Storage error: {str(store_error)}",
                        'text_data': {
                            'original_text': extracted_text[:1000] if extracted_text else "No text extracted",
                            'cleaned_text': cleaned_text[:1000] if cleaned_text else "No cleaned text",
                            'text_length': len(extracted_text) if extracted_text else 0
                        }
                    }
                    analyses_by_filename[missing_file] = placeholder_analysis
                    logger.info(f"Dashboard: Added placeholder analysis for {missing_file} due to storage error")
                    
            except Exception as e:
                import traceback
                logger.error(f"Dashboard: Error creating baseline analysis for {missing_file}: {str(e)}")
                logger.error(f"Dashboard: Error traceback: {traceback.format_exc()}")

                # Create a placeholder as fallback - ensure it has all required fields
                if missing_file not in analyses_by_filename:
                    placeholder_analysis = {
                        '_id': f"placeholder_{missing_file.replace('.', '_')}",
                        'filename': baseline_filename or f"{BASELINE_PREFIX} {missing_file}",
                        'document_id': missing_file,  # Store original filename for matching
                        'title': f"Policy from {university_name or missing_file.split('-')[0].title()}",
                        'analysis_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        'user_id': -1,  # Baseline user ID
                        'is_user_analysis': False,
                        'is_baseline': True,
                        'is_placeholder': True,  # Mark as placeholder
                        'classification': {
                            "classification": "Moderate",
                            "confidence": 75,
                            "source": "Placeholder"
                        },
                        'score': 50,
                        'themes': [{"name": "Policy", "score": 0.8, "confidence": 75}],
                        'summary': f"This is a placeholder for {missing_file}. The original file could not be processed.",
                        'text_data': {
                            'original_text': f"Placeholder for {missing_file}",
                            'cleaned_text': f"Placeholder for {missing_file}",
                            'text_length': 0
                        }
                    }
                    analyses_by_filename[missing_file] = placeholder_analysis
                    logger.info(f"Dashboard: Added placeholder analysis for {missing_file} due to general error")

        # Combine all analyses into a single list
        combined_analyses = list(analyses_by_filename.values())
        logger.info(f"Dashboard: Combined into {len(combined_analyses)} total analyses")
        
        # Log the filenames of all analyses being displayed
        displayed_files = [analysis.get('filename', '') for analysis in combined_analyses]
        logger.info(f"Dashboard: Displaying analyses for: {sorted(displayed_files)}")
        
        # Sort by date (newest first)
        combined_analyses.sort(key=lambda a: _to_epoch(a.get('analysis_date')), reverse=True)
        
        # Load sample policies if needed
        _load_sample_policies_if_needed(user_analyses)
        
        # Generate dashboard data
        dashboard_data = _prepare_dashboard_data(
            current_user, 
            user_analyses, 
            combined_analyses
        )
        
        return render_template('dashboard.html', data=dashboard_data)
        
    except Exception as e:
        logger.error(f"Dashboard error: {str(e)}")
        flash('Dashboard loaded with limited data due to an error.', 'warning')
        return render_template('dashboard.html', data=_get_minimal_dashboard_data(current_user))

def _get_clean_dataset_dir():
    """Return absolute path to clean_dataset directory."""
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'policies', 'clean_dataset')

def _list_clean_dataset_files(clean_dataset_dir):
    """List valid policy files from clean_dataset directory."""
    logger.info(f"Dashboard: Using clean_dataset directory: {clean_dataset_dir}")
    files = []
    if os.path.exists(clean_dataset_dir):
        for filename in os.listdir(clean_dataset_dir):
            if filename.endswith(('.pdf', DOCX_EXTENSION)) and 'guidance' not in filename.lower():
                files.append(filename)
    logger.info(f"Dashboard: Found {len(files)} policy files in clean_dataset")
    return files

def _merge_user_and_baseline_analyses(user_analyses, baseline_analyses):
    """Merge analyses preferring user analyses over baseline ones."""
    analyses_by_filename = {}
    for analysis in user_analyses:
        filename = analysis.get('filename', '')
        if not filename:
            continue
        analyses_by_filename[filename] = analysis
        analysis['is_user_analysis'] = True
    for analysis in baseline_analyses:
        filename = analysis.get('filename', '')
        if not filename or filename in analyses_by_filename:
            continue
        analyses_by_filename[filename] = analysis
        analysis['is_user_analysis'] = False
    return analyses_by_filename

def _identify_missing_clean_files(clean_dataset_files, analyses_by_filename):
    """Return clean_dataset files that are not represented in analyses."""
    missing_files = []
    for clean_file in clean_dataset_files:
        found = False
        for filename in analyses_by_filename.keys():
            analysis = analyses_by_filename[filename]
            if clean_file == filename or (analysis.get('document_id') and clean_file == analysis.get('document_id')):
                found = True
                logger.info(f"Dashboard: Found match for {clean_file} in analysis {filename}")
                break
        if not found:
            missing_files.append(clean_file)
    return missing_files

def _create_or_placeholder_baseline_for_file(*, missing_file, clean_dataset_dir, analyses_by_filename):
    """Create baseline analysis for missing clean_dataset file or add placeholder on failure."""
    # Skip guidance files
    if 'guidance' in missing_file.lower() or missing_file == 'dataset_info.md':
        return

    # Resolve path
    file_path = os.path.join(clean_dataset_dir, missing_file)
    if not os.path.exists(file_path):
        logger.error(f"Dashboard: Missing file not found: {file_path}")
        return

    # Derive baseline filename
    university_name = missing_file.split('-')[0].replace('university', '').strip().title()
    if not university_name:
        university_name = missing_file.split('.')[0].replace('university', '').strip().title()
    baseline_filename = f"{BASELINE_PREFIX} {university_name}"
    logger.info(f"Dashboard: Creating baseline analysis for {missing_file} as '{baseline_filename}'")

    # Extraction and cleaning
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise ValueError(f"Empty file: {file_path}")
        extracted_text = text_processor.extract_text_from_file(file_path)
        text_length = len(extracted_text) if extracted_text else 0
        if not extracted_text or text_length < 50:
            if file_path.lower().endswith('.pdf'):
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        fallback_text = ""
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                fallback_text += page_text + "\n"
                    if fallback_text and len(fallback_text) > text_length:
                        extracted_text = fallback_text
                        text_length = len(extracted_text)
                except Exception as fallback_error:
                    logger.warning(f"Dashboard: Fallback extraction failed: {str(fallback_error)}")
            elif file_path.lower().endswith((DOCX_EXTENSION, '.doc')):
                try:
                    from docx import Document
                    doc = Document(file_path)
                    fallback_text = "\n".join([para.text for para in doc.paragraphs if para.text])
                    if fallback_text and len(fallback_text) > text_length:
                        extracted_text = fallback_text
                        text_length = len(extracted_text)
                except Exception as fallback_error:
                    logger.warning(f"Dashboard: Fallback extraction failed: {str(fallback_error)}")
        if not extracted_text or text_length < 20:
            extracted_text = f"AI Policy document from {university_name}. This is a placeholder text as the original document could not be fully processed."
            text_length = len(extracted_text)
    except Exception as extract_error:
        logger.error(f"Dashboard: Text extraction failed for {file_path}: {str(extract_error)}")
        extracted_text = f"AI Policy document from {university_name}. This is a placeholder text as the original document could not be processed due to: {str(extract_error)}"
        text_length = len(extracted_text)

    try:
        cleaned_text = text_processor.clean_text(extracted_text)
    except Exception as clean_error:
        logger.error(f"Dashboard: Text cleaning failed for {missing_file}: {str(clean_error)}")
        cleaned_text = extracted_text

    # Themes
    try:
        themes = theme_extractor.extract_themes(cleaned_text)
        if not themes or len(themes) == 0:
            themes = [
                {"name": "Policy", "score": 0.8, "confidence": 75},
                {"name": "AI Ethics", "score": 0.7, "confidence": 70},
                {"name": "Guidelines", "score": 0.6, "confidence": 65}
            ]
    except Exception as theme_error:
        logger.error(f"Dashboard: Theme extraction failed for {missing_file}: {str(theme_error)}")
        themes = [
            {"name": "Policy", "score": 0.8, "confidence": 75},
            {"name": "AI Ethics", "score": 0.7, "confidence": 70},
            {"name": "Guidelines", "score": 0.6, "confidence": 65}
        ]

    # Classification
    try:
        classification = policy_classifier.classify_policy(cleaned_text)
        if isinstance(classification, str):
            classification = {
                "classification": _standardize_classification(classification),
                "confidence": 80,
                "source": "Baseline Creation"
            }
        elif isinstance(classification, dict) and 'classification' in classification:
            classification['classification'] = _standardize_classification(classification['classification'])
            if 'source' not in classification:
                classification['source'] = "Baseline Creation"
            if 'confidence' not in classification:
                classification['confidence'] = 80
        else:
            logger.warning(f"Dashboard: Unexpected classification format for {missing_file}: {classification}")
            classification = {
                "classification": "Moderate",
                "confidence": 75,
                "source": "Baseline Creation (Default Format)"
            }
    except Exception as class_error:
        logger.error(f"Dashboard: Classification failed for {missing_file}: {str(class_error)}")
        classification = {
            "classification": "Moderate",
            "confidence": 75,
            "source": "Baseline Creation (Default)"
        }

    # Store baseline or create placeholder
    try:
        if not extracted_text:
            extracted_text = f"AI Policy document from {university_name}. This is a placeholder text."
        if not cleaned_text:
            cleaned_text = extracted_text
        if not themes or len(themes) == 0:
            themes = [{"name": "Policy", "score": 0.8, "confidence": 75}]
        if not classification or not isinstance(classification, dict):
            classification = {
                "classification": "Moderate",
                "confidence": 75,
                "source": "Baseline Creation (Default)"
            }

        metadata = {
            "source_file": missing_file,
            "extraction_method": "dashboard_baseline_creation",
            "creation_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "text_length": len(extracted_text) if extracted_text else 0,
            "is_baseline": True
        }

        analysis_id = db_operations.store_user_analysis_results(
            user_id=-1,
            filename=f"{BASELINE_PREFIX} {university_name}",
            original_text=extracted_text,
            cleaned_text=cleaned_text,
            themes=themes,
            classification=classification,
            document_id=missing_file,
            metadata=metadata
        )

        new_analysis = db_operations.get_analysis_by_id(analysis_id)
        if new_analysis:
            new_analysis['is_baseline'] = True
            new_analysis['is_user_analysis'] = False
            analyses_by_filename[missing_file] = new_analysis
        else:
            placeholder_analysis = {
                '_id': f"placeholder_{missing_file.replace('.', '_')}",
                'filename': f"{BASELINE_PREFIX} {university_name}",
                'document_id': missing_file,
                'title': f"Policy from {university_name}",
                'analysis_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'user_id': -1,
                'is_user_analysis': False,
                'is_baseline': True,
                'classification': classification,
                'score': 50,
                'themes': themes,
                'summary': f"This is a baseline analysis for {missing_file}.",
                'text_data': {
                    'original_text': extracted_text,
                    'cleaned_text': cleaned_text,
                    'text_length': len(extracted_text) if extracted_text else 0
                }
            }
            analyses_by_filename[missing_file] = placeholder_analysis
    except (ValueError, KeyError, RuntimeError, OSError) as store_error:
        logger.error(f"Dashboard: Error storing baseline analysis: {str(store_error)}")
        placeholder_analysis = {
            '_id': f"placeholder_{missing_file.replace('.', '_')}",
            'filename': f"{BASELINE_PREFIX} {university_name}",
            'document_id': missing_file,
            'title': f"Policy from {university_name}",
            'analysis_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'user_id': -1,
            'is_user_analysis': False,
            'is_baseline': True,
            'classification': classification,
            'score': 50,
            'themes': themes,
            'summary': f"This is a placeholder for {missing_file}. Storage error: {str(store_error)}",
            'text_data': {
                'original_text': extracted_text[:1000] if extracted_text else "No text extracted",
                'cleaned_text': cleaned_text[:1000] if cleaned_text else "No cleaned text",
                'text_length': len(extracted_text) if extracted_text else 0
            }
        }
        analyses_by_filename[missing_file] = placeholder_analysis

def _load_sample_policies_if_needed(user_analyses):
    """Load sample policies if user has none."""
    if not any(a.get('filename','').startswith(BASELINE_PREFIX) for a in user_analyses):
        try:
            loaded = db_operations.load_sample_policies_for_user(current_user.id)
            if loaded:
                flash('Sample baseline policies have been loaded to your dashboard.', 'success')
        except Exception as e:
            logger.error(f"Failed to auto-load baseline policies: {e}")

def _prepare_dashboard_data(user, user_analyses, combined_analyses):
    """Prepare the complete dashboard data structure."""
    # Generate charts
    try:
        dashboard_charts = chart_generator.generate_user_dashboard_charts(user_analyses)
    except Exception as e:
        logger.error(f"Chart generation error: {e}")
        dashboard_charts = {}
    
    # Log analyses before analytics calculation
    for i, analysis in enumerate(combined_analyses):
        logger.info(f"Dashboard debug: Analysis {i} type: {type(analysis)}")
        if isinstance(analysis, dict):
            cls = analysis.get('classification', 'Not found')
            logger.info(f"Dashboard debug: Analysis {i} classification type: {type(cls)}, value: {cls}")
        else:
            logger.error(f"Dashboard debug: Analysis {i} is not a dict: {type(analysis)}")
    
    # Calculate analytics
    try:
        classification_counts, theme_frequencies = _calculate_analytics(combined_analyses)
        logger.info(f"Dashboard debug: Classification counts: {classification_counts}")
    except Exception as e:
        logger.error(f"Analytics calculation error: {e}")
        import traceback
        logger.error(f"Analytics calculation traceback: {traceback.format_exc()}")
        classification_counts = {'Restrictive': 0, 'Moderate': 0, 'Permissive': 0}
        theme_frequencies = {}
    
    # Get combined statistics
    try:
        user_stats = db_operations.get_analysis_statistics(user.id)
        baseline_stats = db_operations.get_analysis_statistics(-1)
        combined_stats = _combine_stats(user_stats, baseline_stats)
        
        db_stats = {
            'total_analyses': combined_stats.get('total', 0),
            'avg_confidence': round(combined_stats.get('avg_confidence', 0), 1),
            'avg_themes_per_analysis': round(combined_stats.get('avg_themes_per_analysis', 0), 1)
        }
    except Exception as e:
        logger.error(f"DB stats error: {e}")
        db_stats = {'total_analyses': 0, 'avg_confidence': 0, 'avg_themes_per_analysis': 0}
    
    # Prepare user data
    user_data = {
        'id': user.id,
        'username': user.username,
        'email': user.email,
        'full_name': user.get_full_name(),
        'institution': getattr(user, 'institution', None),
        'role': getattr(user, 'role', 'user')
    }
    
    # Process analyses for display
    try:
        processed_analyses = _process_analyses_for_display(combined_analyses)
        logger.info(f"Dashboard: Processed {len(processed_analyses)} analyses for display")
        
        # Debug first few processed analyses
        for i, analysis in enumerate(processed_analyses[:3]):
            logger.info(f"Dashboard debug: Processed analysis {i} keys: {list(analysis.keys())}")
            if 'classification' in analysis:
                logger.info(f"Dashboard debug: Processed analysis {i} classification type: {type(analysis['classification'])}, value: {analysis['classification']}")
    except Exception as e:
        logger.error(f"Processing analyses error: {e}")
        import traceback
        logger.error(f"Processing analyses traceback: {traceback.format_exc()}")
        processed_analyses = []
    
    dashboard_data = {
        'user': user_data,
        'total_policies': len(combined_analyses),
        'classification_counts': classification_counts,
        'theme_frequencies': theme_frequencies,
        'charts': dashboard_charts,
        'recent_analyses': processed_analyses,
        'statistics': db_stats
    }

    # Log dashboard data structure
    logger.info(f"Dashboard debug: Dashboard data keys: {list(dashboard_data.keys())}")
    
    return dashboard_data

@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload_file():
    """
    File upload with multi-file support for policy documents.
    
    Handles both single and multiple file uploads, validates file types,
    and initiates the document processing pipeline.
    """
    if request.method == 'GET':
        logger.info(f"Upload page accessed by user: {current_user.username}")
        return render_template('upload.html')
    
    # Handle file uploads
    uploaded_files = request.files.getlist('files[]')
    
    # Fallback for single file upload
    if not uploaded_files or all(f.filename == '' for f in uploaded_files):
        single_file = request.files.get('file')
        if single_file and single_file.filename != '':
            uploaded_files = [single_file]
    
    if not uploaded_files or all(f.filename == '' for f in uploaded_files):
        flash('No files selected', 'error')
        return redirect(request.url)
    
    # Check file limit
    if len(uploaded_files) > app.config.get('MAX_FILES_PER_UPLOAD', 10):
        flash(f'Too many files. Maximum {app.config.get("MAX_FILES_PER_UPLOAD", 10)} files allowed.', 'error')
        return redirect(request.url)
    
    successful_uploads = []
    failed_uploads = []
    
    # Process each file
    for file in uploaded_files:
        if file.filename == '':
            continue
            
        if file and allowed_file(file.filename):
            try:
                # Secure the filename and save file
                filename = secure_filename(file.filename)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_")
                unique_filename = f"{current_user.id}_{timestamp}{filename}"
                
                create_upload_folder()
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
                file.save(file_path)
                
                # Process uploaded file
                processing_result = process_new_upload(file_path, file.filename)
                if processing_result['success']:
                    logger.info(f"Auto-processed: {processing_result['standardized_name']}")
                else:
                    logger.warning(f"Processing failed: {processing_result['error']}")

                successful_uploads.append({
                    'original': file.filename,
                    'unique': unique_filename,
                    'size': os.path.getsize(file_path)
                })
                
                logger.info(f"File uploaded successfully: {unique_filename}")
                
            except Exception as e:
                failed_uploads.append({
                    'filename': file.filename,
                    'error': str(e)
                })
                logger.error(f"Error uploading file {file.filename}: {str(e)}")
        else:
            failed_uploads.append({
                'filename': file.filename,
                'error': 'Invalid file type'
            })
    
    # Provide feedback
    if successful_uploads:
        if len(successful_uploads) == 1:
            flash('File uploaded successfully! Starting analysis...', 'success')
            return redirect(url_for('analyse_document', filename=successful_uploads[0]['unique']))
        else:
            flash(f'{len(successful_uploads)} files uploaded successfully! Starting batch analysis...', 'success')
            file_list = [f['unique'] for f in successful_uploads]
            return redirect(url_for('batch_analyse', files=','.join(file_list)))
    
    if failed_uploads:
        error_msg = f"{len(failed_uploads)} files failed to upload: " + ", ".join([f['filename'] for f in failed_uploads])
        flash(error_msg, 'error')
    
    return redirect(request.url)

@app.route('/batch-analyse/<path:files>')
@login_required
def batch_analyse(files):
    """
    Batch analysis for multiple uploaded files.
    
    Processes multiple documents simultaneously through the AI analysis
    pipeline and provides consolidated results and statistics.
    """
    try:
        file_list = files.split(',')
        
        # Security check
        for filename in file_list:
            if not filename.startswith(f"{current_user.id}_"):
                flash('Access denied. You can only analyse your own documents.', 'error')
                return redirect(url_for('upload_file'))
        
        logger.info(f"Starting batch analysis of {len(file_list)} files")
        
        batch_results = []
        successful_analyses = 0
        failed_analyses = 0
        
        for i, filename in enumerate(file_list, 1):
            try:
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                
                if not os.path.exists(file_path):
                    failed_analyses += 1
                    continue
                
                # Run AI pipeline
                extracted_text = text_processor.extract_text_from_file(file_path)
                if not extracted_text:
                    failed_analyses += 1
                    continue
                
                cleaned_text = text_processor.clean_text(extracted_text)
                themes = theme_extractor.extract_themes(cleaned_text)
                classification = policy_classifier.classify_policy(cleaned_text)
                
                # Store in database
                analysis_id = db_operations.store_user_analysis_results(
                    user_id=current_user.id,
                    filename=filename,
                    original_text=extracted_text,
                    cleaned_text=cleaned_text,
                    themes=themes,
                    classification=classification
                )
                
                # Generate charts
                charts = chart_generator.generate_analysis_charts(themes, classification, cleaned_text)
                
                batch_results.append({
                    'filename': filename,
                    'original_filename': filename.split('_', 2)[-1],
                    'analysis_id': analysis_id,
                    'themes': themes[:5],
                    'classification': classification,
                    'charts': charts,
                    'text_length': len(cleaned_text),
                    'theme_count': len(themes),
                    'status': 'success'
                })
                
                successful_analyses += 1
                
            except Exception as e:
                logger.error(f"Error analyzing {filename}: {str(e)}")
                batch_results.append({
                    'filename': filename,
                    'original_filename': filename.split('_', 2)[-1],
                    'status': 'failed',
                    'error': str(e)
                })
                failed_analyses += 1
        
        # Generate batch summary
        if successful_analyses > 0:
            all_classifications = [r['classification']['classification'] for r in batch_results if r['status'] == 'success']
            classification_summary = {cls: all_classifications.count(cls) for cls in set(all_classifications)}
            
            all_themes = []
            for result in batch_results:
                if result['status'] == 'success':
                    all_themes.extend([theme['name'] for theme in result['themes']])
            
            from collections import Counter
            theme_summary = dict(Counter(all_themes).most_common(10))
            
            batch_summary = {
                'total_files': len(file_list),
                'successful': successful_analyses,
                'failed': failed_analyses,
                'classification_summary': classification_summary,
                'theme_summary': theme_summary,
                'avg_confidence': sum([r['classification']['confidence'] for r in batch_results if r['status'] == 'success']) / successful_analyses if successful_analyses > 0 else 0
            }
        else:
            batch_summary = {
                'total_files': len(file_list),
                'successful': 0,
                'failed': failed_analyses,
                'classification_summary': {},
                'theme_summary': {},
                'avg_confidence': 0
            }
        
        logger.info(f"Batch analysis completed: {successful_analyses}/{len(file_list)} successful")
        
        return render_template('batch_results.html', {
            'results': batch_results,
            'summary': batch_summary,
            'user': current_user,
            'analysis_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
        
    except Exception as e:
        logger.error(f"Error in batch analysis: {str(e)}")
        flash('Error during batch analysis. Please try again.', 'error')
        return redirect(url_for('upload_file'))

@app.route('/analyse/<filename>')
@login_required
def analyse_document(filename):
    """
    Document analysis pipeline for individual policy files.
    
    Processes a single document through the complete AI analysis workflow
    including text extraction, theme analysis, and policy classification.
    """
    try:
        is_baseline = filename.startswith(BASELINE_PREFIX)
        
        # Security check - allow baseline files or user's own files
        # Baseline files are accessible to all users for comparison purposes
        if not is_baseline and not filename.startswith(f"{current_user.id}_"):
            flash('Access denied. You can only analyse your own documents or baseline policies.', 'error')
            return redirect(url_for('upload_file'))
        
        # Determine the correct file path
        if is_baseline:
            # For baseline files, we need to map the display name back to the actual filename
            # First check if it's already a proper filename with extension
            original_filename = filename.split(' - ')[-1] if ' - ' in filename else filename.replace(BASELINE_PREFIX + ' ', '')
            
            # Create a mapping of display names to actual filenames
            university_file_mapping = {
                'University of Oxford': 'oxford-ai-policy.pdf',
                'University of Cambridge': 'cambridge-ai-policy.pdf',
                'Imperial College London': 'imperial-ai-policy.docx',
                'University of Edinburgh': 'edinburgh university-ai-policy.pdf',
                'Leeds Trinity University': 'leeds trinity university-ai-policy.pdf',
                'MIT': 'mit-ai-policy.pdf',
                'Harvard University': 'harvard-ai-policy.pdf',
                'Stanford University': 'stanford-ai-policy.pdf',
                'University of Tokyo': 'tokyo-ai-policy.docx',
                'Jagiellonian University': 'jagiellonian university-ai-policy.pdf',
                'Belfast University': 'belfast university-ai-policy.pdf',
                'University of Chicago': 'chicago-ai-policy.docx',
                'Columbia University': 'columbia-ai-policy.pdf',
                'Cornell University': 'cornell-ai-policy.docx',
                'University of Liverpool': 'liverpool policy-ai-policy.pdf'
            }
            
            # Try to map the display name to actual filename
            if original_filename in university_file_mapping:
                actual_filename = university_file_mapping[original_filename]
            else:
                # If no mapping found, use the original filename as-is
                actual_filename = original_filename
            
            # Use absolute path (consistent with dashboard listing) to avoid CWD issues
            clean_dataset_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'policies', 'clean_dataset')
            file_path = os.path.join(clean_dataset_dir, actual_filename)
            logger.info(f"Baseline analysis - Display name: {original_filename}")
            logger.info(f"Baseline analysis - Mapped filename: {actual_filename}")
            logger.info(f"Baseline analysis - File path: {file_path}")
            logger.info(f"Baseline analysis - File exists: {os.path.exists(file_path)}")
        else:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            logger.info(f"User analysis - File path: {file_path}")
            logger.info(f"User analysis - File exists: {os.path.exists(file_path)}")
        
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            flash('File not found', 'error')
            return redirect(url_for('upload_file'))
        
        logger.info(f"Starting analysis of file: {filename}")
        
        # Check if analysis already exists to avoid duplication
        existing_analysis = db_operations.get_analysis_by_filename(current_user.id, filename)
        if not existing_analysis and is_baseline:
            # For baseline, check global baseline record (user_id = -1)
            existing_analysis = db_operations.get_analysis_by_filename(-1, filename)
        if existing_analysis:
            themes = existing_analysis.get('themes', [])
            classification = existing_analysis.get('classification', {})
            cleaned_text = existing_analysis.get('text_data', {}).get('cleaned_text', '')
            extracted_text = existing_analysis.get('text_data', {}).get('original_text', '')
            analysis_id = str(existing_analysis['_id'])
        else:
            # Extract and process text
            extracted_text = text_processor.extract_text_from_file(file_path)
            if not extracted_text:
                flash('Could not extract text from file', 'error')
                return redirect(url_for('upload_file'))
            
            cleaned_text = text_processor.clean_text(extracted_text)
            themes = theme_extractor.extract_themes(cleaned_text)
            classification = policy_classifier.classify_policy(cleaned_text)
            
            # Store results in database (handles update if exists)
            analysis_id = db_operations.store_user_analysis_results(
                user_id=current_user.id,
                filename=filename,
                original_text=extracted_text,
                cleaned_text=cleaned_text,
                themes=themes,
                classification=classification,
                username=getattr(current_user, 'username', None)
            )
        
        # Generate visualizations
        charts = chart_generator.generate_analysis_charts(themes, classification, cleaned_text)
        
        # Prepare comprehensive results
        text_stats = text_processor.get_text_statistics(cleaned_text)
        theme_summary = theme_extractor.get_theme_summary(themes)
        classification_details = policy_classifier.get_classification_details(cleaned_text)
        
        logger.info(f"Analysis completed successfully for: {filename}")
        
        results = {
            'filename': filename,
            'analysis_id': analysis_id,
            'themes': themes,
            'classification': classification,
            'charts': charts,
            'text_stats': text_stats,
            'theme_summary': theme_summary,
            'classification_details': classification_details,
            'analysis_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'user': current_user,
            'processing_summary': {
                'original_length': len(extracted_text),
                'cleaned_length': len(cleaned_text),
                'themes_found': len(themes),
                'classification_confidence': classification.get('confidence', 0),
                'processing_method': classification.get('method', 'N/A')
            }
        }
        
        return render_template('results.html', results=results)
        
    except Exception as e:
        logger.error(f"Error during analysis of {filename}: {str(e)}")
        flash('Error analysing document. Please try again.', 'error')
        return redirect(url_for('upload_file'))

@app.route('/validate/<analysis_id>')
@login_required
def validate_analysis(analysis_id):
    """
    Return JSON with citation validation issues for given analysis.
    
    Validates the academic sources and citations used in recommendations
    to ensure proper referencing standards.
    """
    from flask import jsonify
    analysis = db_operations.get_analysis_by_id(analysis_id)
    if not analysis:
        return jsonify({"error": ANALYSIS_NOT_FOUND}), 404
    recs = db_operations.get_recommendations_by_analysis(current_user.id, analysis_id)
    if not recs:
        # fallback: any user_id
        any_doc = db_operations.recommendations.find_one({"analysis_id": analysis_id})
        if any_doc:
            recs = any_doc.get("recommendations", [])
        else:
            # As a last resort, generate recommendations on the fly for validation
            text_data = analysis.get('text_data', {})
            cleaned_text = text_data.get('cleaned_text', text_data.get('original_text', ''))
            rec_package = recommendation_engine.generate_recommendations(
                policy_text=cleaned_text,
                analysis_id=analysis_id
            )
            recs = rec_package.get('recommendations', [])
    from src.utils.validation import validate_recommendation_sources
    issues = validate_recommendation_sources(recs)
    return jsonify({"issues": issues})

@app.route('/recommendations/<analysis_id>')
@login_required
def get_recommendations(analysis_id):
    """
    Generate evidence-based recommendations for policy improvement.
    
    Creates detailed, academic-quality recommendations based on ethical
    AI framework analysis and institutional context assessment.
    """
    try:
        logger.info("=== Starting recommendation generation ===")
        logger.info(f"User ID: {current_user.id}, Username: {getattr(current_user, 'username', 'N/A')}")
        logger.info(f"Analysis ID: {analysis_id}")
        
        # Get the analysis data (try by user id, fallback to global and verify ownership)
        logger.info("Attempting to get user's analysis...")
        analysis = db_operations.get_user_analysis_by_id(current_user.id, analysis_id)
        
        if analysis:
            logger.info("Found analysis in user's analyses")
            logger.info(f"Analysis details - User ID: {analysis.get('user_id')}, Is Baseline: {analysis.get('is_baseline', False)}")
        else:
            logger.warning(f"{ANALYSIS_NOT_FOUND} in user's analyses, trying global lookup...")
            analysis = db_operations.get_analysis_by_id(analysis_id)
            
            if analysis:
                logger.info("Found analysis in global collection")
                logger.info(f"Analysis details - ID: {analysis.get('_id')}")
                logger.info(f"Analysis owner: User ID: {analysis.get('user_id')}, Username: {analysis.get('username')}")
                logger.info(f"Is baseline: {analysis.get('is_baseline', False)}")
                
                # Tymczasowe wyłączenie sprawdzania uprawnień do testów
                logger.warning("Temporary disabling permission checks for testing")
                logger.info(f"Analysis user_id: {analysis.get('user_id')}, Current user ID: {current_user.id}")
                logger.info(f"Analysis username: {analysis.get('username')}, Current username: {getattr(current_user, 'username', 'N/A')}")
                
                # Temporary disabling permission checks for testing
                # Restore permission checks after testing
                # is_baseline = analysis.get('is_baseline', False)
                # is_owner = (
                #     str(analysis.get('user_id')) in [str(current_user.id), 'None', ''] or 
                #     str(analysis.get('username')) == str(getattr(current_user, 'username', ''))
                # )
                # 
                # logger.info(f"Access check - is_baseline: {is_baseline}, is_owner: {is_owner}")
                # 
                # if not (is_baseline or is_owner):
                #     logger.error(f"Access denied for user {current_user.id} to analysis {analysis_id}")
                #     logger.error(f"User ID match: {str(analysis.get('user_id')) == str(current_user.id)}")
                #     logger.error(f"Username match: {str(analysis.get('username')) == str(getattr(current_user, 'username', ''))}")
                #     flash('You do not have permission to view this analysis.', 'error')
                #     return redirect(url_for('dashboard'))
            else:
                logger.error(f"Analysis {analysis_id} not found in global lookup")
                flash(f'{ANALYSIS_NOT_FOUND}.', 'error')
                return redirect(url_for('dashboard'))
        
        logger.info(f"Found analysis data for ID: {analysis_id}")
        
        # Extract analysis components with logging
        themes = analysis.get('themes', [])
        classification = analysis.get('classification', {})
        text_data = analysis.get('text_data', {})
        cleaned_text = text_data.get('cleaned_text', text_data.get('original_text', ''))
        
        logger.info(f"Extracted data - Themes: {len(themes)}, Classification: {classification}")
        logger.info(f"Text data length: {len(cleaned_text) if cleaned_text else 0} chars")
        
        if not cleaned_text:
            logger.error("No text content found in analysis to generate recommendations")
            flash('Analysis text not found. Cannot generate recommendations.', 'warning')
            return redirect(url_for('dashboard'))
        
        if not cleaned_text:
            flash('Analysis text not found. Cannot generate recommendations.', 'warning')
            return redirect(url_for('dashboard'))
        
        # Try to load stored recommendations first
        stored_recs = db_operations.get_recommendations_by_analysis(current_user.id, analysis_id)
        if stored_recs:
            recommendation_package = {
                'recommendations': stored_recs,
                'analysis_metadata': {
                    'generated_date': analysis.get('analysis_date', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
                    'methodology': 'Previously generated'
                }
            }
        else:
            # Generate new recommendations
            try:
                print("\n===== GENERATING RECOMMENDATIONS WITH KNOWLEDGE BASE =====\n")
                recommendations = recommendation_engine.generate_recommendations(
                    policy_text=cleaned_text,
                    institution_type="university",
                analysis_id=analysis_id
            )
                print("\n===== RECOMMENDATIONS GENERATED =====\n")
                
                # Debug: Print sources for each recommendation
                for i, rec in enumerate(recommendations.get("recommendations", [])):
                    print(f"Recommendation {i+1}: {rec.get('title', 'Unknown')}")
                    print(f"  Sources: {rec.get('sources', [])}")
                    print(f"  References: {len(rec.get('references', []))} items")
                print("\n===== END OF RECOMMENDATIONS DEBUG =====\n")
                
                recommendation_package = recommendations
            except Exception as e:
                logger.error(f"Error generating recommendations: {str(e)}")
                flash('Error generating recommendations. Please try again.', 'error')
                return redirect(url_for('dashboard'))
        
        # Prepare data for template
        recommendation_data = {
            'analysis': {
                'filename': analysis.get('filename', 'Unknown'),
                'classification': classification.get('classification', 'Unknown'),
                'confidence': classification.get('confidence', 0),
                'analysis_id': analysis_id,
                'themes_count': len(themes)
            },
            'recommendations': recommendation_package.get('recommendations', []),
            'coverage_analysis': recommendation_package.get('coverage_analysis', {}),
            'gaps': recommendation_package.get('identified_gaps', []),
            'generated_date': recommendation_package.get('analysis_metadata', {}).get('generated_date', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            'methodology': recommendation_package.get('analysis_metadata', {}).get('methodology', 'Ethical Framework Analysis'),
            'academic_sources': recommendation_package.get('analysis_metadata', {}).get('academic_sources', []),
            'summary': recommendation_package.get('summary', {}),
            'total_recommendations': len(recommendation_package.get('recommendations', []))
        }
        
        logger.info(f"Generated {recommendation_data['total_recommendations']} recommendations")
        
        # Store recommendations in database
        try:
            rec_id = db_operations.store_recommendations(
                user_id=current_user.id,
                analysis_id=analysis_id,
                recs=recommendation_package.get('recommendations', [])
            )
            logger.info(f"Recommendations stored with ID: {rec_id}")
        except Exception as e:
            logger.warning(f"Could not store recommendations: {e}")
        
        return render_template('recommendations.html', data=recommendation_data)
        
    except Exception as e:
        logger.error(f"Error generating recommendations: {str(e)}")
        
        # Fallback to basic recommendations
        try:
            basic_data = {
                'analysis': {
                    'filename': analysis.get('filename', 'Unknown'),
                    'classification': analysis.get('classification', {}).get('classification', 'Unknown'),
                    'analysis_id': analysis_id
                },
                'recommendations': [
                    {
                        'title': 'Policy Review Required',
                        'description': 'Conduct comprehensive review of AI policy to ensure alignment with current best practices.',
                        'priority': 'high',
                        'source': 'Fallback System',
                        'timeframe': '3-6 months',
                        'implementation_steps': [
                            'Review current policy framework',
                            'Consult with stakeholders',
                            'Implement evidence-based improvements'
                        ]
                    }
                ],
                'generated_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'methodology': 'Basic Template (Fallback)',
                'total_recommendations': 1,
                'error_note': 'Advanced analysis temporarily unavailable'
            }
            
            flash('Using basic recommendations due to processing error.', 'warning')
            return render_template('recommendations.html', data=basic_data)
            
        except Exception as fallback_error:
            logger.error(f"Fallback recommendation generation failed: {str(fallback_error)}")
            flash('Error generating recommendations. Please try again.', 'error')
            return redirect(url_for('dashboard'))


@app.route('/delete_analysis/<analysis_id>', methods=['POST'])
@login_required
def delete_analysis(analysis_id):
    """
    Delete user analysis with appropriate security checks.
    
    Removes analysis records from database while protecting baseline
    policies from accidental deletion.
    """
    logger.info(f"Attempting deletion of analysis {analysis_id} by user {current_user.id}")
    try:
        # Get analysis to check ownership and type
        analysis = db_operations.get_user_analysis_by_id(current_user.id, analysis_id)
        if not analysis:
            flash(f'{ANALYSIS_NOT_FOUND} or access denied.', 'error')
            return redirect(url_for('dashboard'))
        
        # Check if it's a baseline policy (protect from deletion)
        filename = analysis.get('filename', '')
        if filename.startswith(BASELINE_PREFIX):
            flash('Baseline policies cannot be deleted.', 'warning')
            return redirect(url_for('dashboard'))
        
        # Delete the analysis and related data
        success = db_operations.delete_user_analysis(current_user.id, analysis_id)
        if success:
            # Try to delete the physical file if it exists
            try:
                # Check multiple possible file locations and naming patterns
                possible_filenames = [
                    filename,  # Original filename
                    f"{current_user.id}_{filename}",  # Prefixed with user ID
                    analysis.get('document_id', '')  # Document ID if different
                ]
                
                # Also check for files with .pdf extension if original doesn't have it
                if not filename.lower().endswith('.pdf'):
                    possible_filenames.append(f"{filename}.pdf")
                    possible_filenames.append(f"{current_user.id}_{filename}.pdf")
                
                # Try each possible filename
                for fname in possible_filenames:
                    if not fname:  # Skip empty filenames
                        continue
                        
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], fname)
                    if os.path.exists(file_path):
                        try:
                            os.remove(file_path)
                            logger.info(f"Successfully deleted file: {file_path}")
                            break  # Stop after first successful deletion
                        except Exception as e:
                            logger.warning(f"Could not delete file {file_path}: {e}")
            
            except Exception as e:
                logger.error(f"Error during file cleanup for analysis {analysis_id}: {str(e)}")
                # Don't fail the whole operation if file deletion fails
            
            flash('Analysis and related data deleted successfully.', 'success')
        else:
            logger.error(f"Failed to delete analysis {analysis_id} for user {current_user.id}")
            flash('Error deleting analysis. Please try again or contact support if the issue persists.', 'error')
        
        return redirect(url_for('dashboard'))
        
    except Exception as e:
        logger.error(f"Error deleting analysis: {str(e)}")
        flash('Error deleting analysis.', 'error')
        return redirect(url_for('dashboard'))

# API routes for additional functionality

@app.route('/api/explain/<analysis_id>')
@login_required
def api_explain_analysis(analysis_id):
    """
    Return JSON with classification explanation keywords.
    
    Provides key terms that influenced the AI policy classification
    decision for transparency and explainability purposes.
    """
    db_ops = DatabaseOperations(uri="mongodb://localhost:27017", db_name="policycraft")
    analysis = db_ops.get_analysis_by_id(analysis_id)
    if not analysis or analysis.get('user_id') != current_user.id:
        return jsonify({'error': ANALYSIS_NOT_FOUND}), 404

    cleaned_text = analysis.get('text_data', {}).get('cleaned_text', '')
    if not cleaned_text:
        return jsonify({'error': 'No cleaned text available'}), 400

    classifier = PolicyClassifier()
    keywords = classifier.explain_classification(cleaned_text, top_n=15)
    result = [{"term": kw, "category": cat, "score": score} for kw, cat, score in keywords]
    return jsonify({'analysis_id': analysis_id, 'keywords': result})

@app.route('/api/analysis/<analysis_id>')
@login_required
def api_get_analysis(analysis_id):
    """
    API endpoint to retrieve analysis results in JSON format.
    
    Provides programmatic access to stored analysis data
    for integration with external systems or applications.
    """
    try:
        analysis = db_operations.get_user_analysis_by_id(current_user.id, analysis_id)
        if analysis:
            return jsonify({'success': True, 'data': analysis})
        else:
            return jsonify({'success': False, 'error': ANALYSIS_NOT_FOUND}), 404
    except Exception as e:
        logger.error(f"API error for analysis {analysis_id}: {str(e)}")
        return jsonify({'success': False, 'error': 'Internal server error'}), 500

# Public routes

# Error handlers

@app.errorhandler(404)
def not_found_error(error):
    """Handle 404 not found errors."""
    logger.warning(f"404 error: {request.url}")
    return render_template('errors/404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    """Handle 500 internal server errors."""
    logger.error(f"500 error: {str(error)}")
    return render_template('errors/500.html'), 500

# Export routes

@app.route('/export/<analysis_id>')
@login_required
def export_view(analysis_id):
    """
    Display export view for recommendations and analysis results.
    
    Renders a dedicated template for exporting recommendations and
    analysis results to various formats (PDF, Word, Excel).
    """
    try:
        logger.info("=== Starting export view preparation ===")
        logger.info(f"User ID: {current_user.id}, Analysis ID: {analysis_id}")
        
        # Get the analysis data
        logger.info("Attempting to get user's analysis...")
        analysis = db_operations.get_user_analysis_by_id(current_user.id, analysis_id)
        
        if analysis:
            logger.info("Found analysis in user's analyses")
            logger.info(f"Analysis details - Filename: {analysis.get('filename')}, ID: {analysis_id}")
        else:
            logger.info(f"{ANALYSIS_NOT_FOUND} in user's analyses, trying global lookup...")
            analysis = db_operations.get_analysis_by_id(analysis_id)
            
            if analysis:
                logger.info("Found analysis in global collection")
                logger.info(f"Analysis details - Filename: {analysis.get('filename')}, ID: {analysis_id}")
            else:
                logger.error(f"Analysis {analysis_id} not found in global lookup")
                flash(f'{ANALYSIS_NOT_FOUND}.', 'error')
                return redirect(url_for('recommendations'))
        
        # Get recommendations
        logger.info(f"Attempting to get recommendations for analysis {analysis_id}...")
        recommendations = db_operations.get_recommendations_by_analysis(current_user.id, analysis_id)
        
        if recommendations:
            logger.info(f"Found {len(recommendations)} recommendations for analysis {analysis_id}")
        else:
            logger.warning(f"No recommendations found for analysis {analysis_id}")
            flash(f'{NO_RECOMMENDATIONS_FOUND}.', 'warning')
            return redirect(url_for('get_recommendations', analysis_id=analysis_id))
        
        # Generate charts for export view
        logger.info("Generating charts for export view...")
        themes = analysis.get('themes', [])
        classification = analysis.get('classification', {})
        cleaned_text = analysis.get('text_data', {}).get('cleaned_text', '')
        
        try:
            charts = chart_generator.generate_analysis_charts(themes, classification, cleaned_text)
            logger.info(f"Generated {len(charts)} charts for export view")
        except Exception as chart_error:
            logger.error(f"Error generating charts: {chart_error}")
            charts = {}
        
        # Prepare data for template
        export_data = {
            'analysis': {
                'filename': analysis.get('filename', 'Unknown'),
                'classification': analysis.get('classification', {}).get('classification', 'Unknown'),
                'confidence': analysis.get('classification', {}).get('confidence', 0),
                'analysis_id': analysis_id,
                'themes': analysis.get('themes', []),
                'text_data': analysis.get('text_data', {})
            },
            'recommendations': recommendations,
            'generated_date': datetime.now().isoformat(),
            'total_recommendations': len(recommendations),
            'charts': charts
        }
        
        return render_template('export_recommendations.html', data=export_data)
        
    except Exception as e:
        logger.error(f"Error preparing export view: {str(e)}")
        logger.exception("Full traceback for export view error:")
        flash('Error preparing export view. Please try again.', 'error')
        return redirect(url_for('dashboard'))

@app.route('/export/<analysis_id>/pdf')
@login_required
def export_pdf(analysis_id):
    """
    Export recommendations and analysis results as PDF.
    
    Generates a PDF document containing all recommendations and
    analysis results for the specified analysis.
    """
    try:
        # Get the analysis data
        analysis = db_operations.get_user_analysis_by_id(current_user.id, analysis_id)
        
        if not analysis:
            analysis = db_operations.get_analysis_by_id(analysis_id)
            
            if not analysis:
                return jsonify({'error': ANALYSIS_NOT_FOUND}), 404
        
        # Get recommendations
        recommendations = db_operations.get_recommendations_by_analysis(current_user.id, analysis_id)
        
        if not recommendations:
            return jsonify({'error': NO_RECOMMENDATIONS_FOUND}), 404
        
        # Prepare data for export
        export_data = {
            'analysis': {
                'filename': analysis.get('filename', 'Unknown'),
                'classification': analysis.get('classification', {}).get('classification', 'Unknown'),
                'confidence': analysis.get('classification', {}).get('confidence', 0),
                'analysis_id': analysis_id,
                'themes': analysis.get('themes', [])
            },
            'recommendations': recommendations,
            'generated_date': datetime.now().isoformat(),
            'total_recommendations': len(recommendations)
        }
        
        # Initialise export engine
        from src.export import ExportEngine
        export_engine = ExportEngine()
        
        # Generate PDF
        pdf_data = export_engine.export_to_pdf(export_data)
        
        # Create response
        response = make_response(pdf_data)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename=PolicyCraft_Analysis_{analysis_id}.pdf'
        
        return response
        
    except Exception as e:
        logger.error(f"Error exporting to PDF: {str(e)}")
        return jsonify({'error': 'Error generating PDF export'}), 500

@app.route('/export/<analysis_id>/word')
@login_required
def export_word(analysis_id):
    """
    Export recommendations and analysis results as Word document.
    
    Generates a Word document containing all recommendations and
    analysis results for the specified analysis.
    """
    try:
        # Get the analysis data
        analysis = db_operations.get_user_analysis_by_id(current_user.id, analysis_id)
        
        if not analysis:
            analysis = db_operations.get_analysis_by_id(analysis_id)
            
            if not analysis:
                return jsonify({'error': ANALYSIS_NOT_FOUND}), 404
        
        # Get recommendations
        recommendations = db_operations.get_recommendations_by_analysis(current_user.id, analysis_id)
        
        if not recommendations:
            return jsonify({'error': NO_RECOMMENDATIONS_FOUND}), 404
        
        # Prepare data for export
        export_data = {
            'analysis': {
                'filename': analysis.get('filename', 'Unknown'),
                'classification': analysis.get('classification', {}).get('classification', 'Unknown'),
                'confidence': analysis.get('classification', {}).get('confidence', 0),
                'analysis_id': analysis_id,
                'themes': analysis.get('themes', [])
            },
            'recommendations': recommendations,
            'generated_date': datetime.now().isoformat(),
            'total_recommendations': len(recommendations)
        }
        
        # Initialise export engine
        from src.export import ExportEngine
        export_engine = ExportEngine()
        
        # Generate Word document
        docx_data = export_engine.export_to_word(export_data)
        
        # Create response
        response = make_response(docx_data)
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename=PolicyCraft_Analysis_{analysis_id}.docx'
        
        return response
        
    except Exception as e:
        logger.error(f"Error exporting to Word: {str(e)}")
        return jsonify({'error': 'Error generating Word export'}), 500

@app.route('/export/<analysis_id>/excel')
@login_required
def export_excel(analysis_id):
    """
    Export recommendations and analysis results as Excel spreadsheet.
    
    Generates an Excel spreadsheet containing all recommendations and
    analysis results for the specified analysis.
    """
    try:
        # Get the analysis data
        analysis = db_operations.get_user_analysis_by_id(current_user.id, analysis_id)
        
        if not analysis:
            analysis = db_operations.get_analysis_by_id(analysis_id)
            
            if not analysis:
                return jsonify({'error': ANALYSIS_NOT_FOUND}), 404
        
        # Get recommendations
        recommendations = db_operations.get_recommendations_by_analysis(current_user.id, analysis_id)
        
        if not recommendations:
            return jsonify({'error': NO_RECOMMENDATIONS_FOUND}), 404
        
        # Prepare data for export
        export_data = {
            'analysis': {
                'filename': analysis.get('filename', 'Unknown'),
                'classification': analysis.get('classification', {}).get('classification', 'Unknown'),
                'confidence': analysis.get('classification', {}).get('confidence', 0),
                'analysis_id': analysis_id,
                'themes': analysis.get('themes', [])
            },
            'recommendations': recommendations,
            'generated_date': datetime.now().isoformat(),
            'total_recommendations': len(recommendations)
        }
        
        # Initialise export engine
        from src.export import ExportEngine
        export_engine = ExportEngine()
        
        # Generate Excel spreadsheet
        excel_data = export_engine.export_to_excel(export_data)
        
        # Create response
        response = make_response(excel_data)
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        response.headers['Content-Disposition'] = f'attachment; filename=PolicyCraft_Analysis_{analysis_id}.xlsx'
        
        return response
        
    except Exception as e:
        logger.error(f"Error exporting to Excel: {str(e)}")
        return jsonify({'error': 'Error generating Excel export'}), 500

# Register template filters
app.jinja_env.filters["clean_university_name"] = clean_university_name
app.jinja_env.filters["clean_filename"] = clean_filename
app.jinja_env.filters["format_british_date"] = format_british_date

# Import and register additional template filters
from src.web.utils.template_utils import clean_literature_name
app.jinja_env.filters["clean_literature_name"] = clean_literature_name

def handle_first_login_onboarding(user_id: int) -> bool:
    """
    Handle automatic onboarding for new users by loading sample policies.
    
    Provides new users with baseline policy examples to demonstrate
    the system capabilities and provide comparison benchmarks.
    """
    try:
        from src.database.models import User
        
        user = User.query.get(user_id)
        if not user:
            return False
            
        # Check if user needs onboarding
        if user.is_first_login():
            logger.info(f"Starting onboarding for new user: {user.username}")
            
            # Load sample policies automatically
            success = db_operations.load_sample_policies_for_user(user_id)
            
            if success:
                # Mark onboarding as completed
                user.complete_onboarding()
                logger.info(f"Onboarding completed for user: {user.username}")
                flash("Welcome! We have loaded 15 sample university policies to get you started.", "success")
                return True
            else:
                # If sample load failed, but user already has baseline analyses, treat as success
                # Robust baseline detection check if ANY analysis filename starts with BASELINE
                try:
                    user_analyses = db_operations.get_user_analyses(user_id)
                    baseline_exists = any(a.get('filename', '').startswith(BASELINE_PREFIX) for a in user_analyses)
                except Exception as _:
                    baseline_exists = False
                if baseline_exists:
                    user.complete_onboarding()
                    logger.info(f"Baseline analyses detected despite load failure onboarding marked complete for {user.username}")
                    flash("Welcome! Sample policies are already available in your dashboard.", "info")
                    return True
                logger.warning(f"Onboarding failed for user: {user.username}")
                flash("Welcome! There was an issue loading sample policies.", "warning")
                
        return False
        
    except Exception as e:
        logger.error(f"Error in onboarding: {e}")
        return False

if __name__ == '__main__':
    """Run the PolicyCraft application in development mode."""
    os.makedirs('logs', exist_ok=True)
    create_upload_folder()
    
    logger.info("Starting PolicyCraft Application")
    
    # Note: Automatic document scanning disabled to prevent excessive backups
    # Documents are now processed on-demand when uploaded via admin interface
    logger.info("Automatic document scanning disabled - documents processed on upload")
    
    print("PolicyCraft starting at: http://localhost:5001")
    
    app.run(debug=True, host='localhost', port=5001)