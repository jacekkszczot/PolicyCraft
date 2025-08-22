"""
Export engine for PolicyCraft - handles exporting recommendations and analysis results
to various formats (PDF, Word, Excel).

Author: Jacek Robert Kszczot
Project: MSc Data Science & AI - COM7016
University: Leeds Trinity University
"""

import os
import io
import json
import base64
import logging
import re
import unicodedata
from datetime import datetime
from typing import Dict, Any, List, Optional, Union, Tuple

# PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, ListFlowable, ListItem
from reportlab.lib.units import inch

# Word document generation
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Excel generation
import xlsxwriter

logger = logging.getLogger(__name__)

class ExportEngine:
    """
    Handles exporting recommendations and analysis results to various formats.
    """
    
    def __init__(self, export_dir: str = "exports"):
        """
        Initialise the export engine.
        
        Args:
            export_dir: Directory to store exported files
        """
        self.export_dir = export_dir
        os.makedirs(export_dir, exist_ok=True)
        
        # Define common styles
        self.title_colour = "#3498db"  # Blue
        self.heading_colour = "#2c3e50"  # Dark blue
        self.accent_colour = "#e74c3c"  # Red
        self.text_colour = "#34495e"  # Dark grey
        
        # Chart image settings
        self.chart_width = 600
        self.chart_height = 400
        
    def _format_date(self, date_str: str) -> str:
        """Format date string to British DD/MM/YYYY style."""
        try:
            if isinstance(date_str, str):
                if "T" in date_str:
                    # ISO format
                    dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
                    return dt.strftime("%d/%m/%Y")
                elif "-" in date_str:
                    # YYYY-MM-DD format
                    parts = date_str.split("-")
                    if len(parts) == 3:
                        return f"{parts[2]}/{parts[1]}/{parts[0]}"
            return date_str
        except Exception:
            return date_str

    def _clean_text(self, txt: str) -> str:
        """Normalise and replace problematic unicode characters across exports."""
        if not isinstance(txt, str):
            return txt
        t = unicodedata.normalize('NFKC', txt)
        # Non-breaking spaces and variations -> regular space
        t = t.replace('\u00A0', ' ').replace('\u202F', ' ').replace('\u2007', ' ').replace('\u2009', ' ')
        # Soft hyphen & non-breaking hyphen -> '-'
        t = t.replace('\u00AD', '').replace('\u2011', '-').replace('\u2010', '-')
        # En/em dashes -> '-'
        t = t.replace('\u2013', '-').replace('\u2014', '-')
        # Bullet/black square -> '- '
        t = t.replace('•', '- ').replace('■', '')
        # Collapse whitespace
        t = re.sub(r"[\t\x0b\x0c\r]+", " ", t)
        t = re.sub(r"\s{2,}", " ", t)
        return t.strip()
            
    def _convert_chart_to_image(self, chart_data: Dict) -> Optional[bytes]:
        """
        Convert a Plotly chart to an image.
        
        Args:
            chart_data: Dictionary containing Plotly chart data and layout
            
        Returns:
            Image as bytes or None if conversion fails
        """
        try:
            # Ensure static image engine is available
            try:
                import kaleido  # noqa: F401  # required for plotly static image export
            except Exception:
                logger.warning(
                    "Plotly static image export requires 'kaleido'. "
                    "Charts will be skipped. Install with: pip install -U kaleido"
                )
                return None
            # Import plotly here to avoid circular imports
            import plotly.graph_objects as go
            import plotly.io as pio
            
            # Create a figure from the chart data
            if isinstance(chart_data, dict) and 'data' in chart_data and 'layout' in chart_data:
                fig = go.Figure(data=chart_data['data'], layout=chart_data['layout'])
                
                # Set the size of the image
                fig.update_layout(
                    width=self.chart_width,
                    height=self.chart_height
                )
                
                # Convert to PNG image
                img_bytes = pio.to_image(fig, format='png')
                return img_bytes
            else:
                logger.warning("Invalid chart data format: %s", type(chart_data))
                return None
        except Exception as e:
            logger.warning("Error converting chart to image: %s", str(e))
            return None
            
    def _process_charts_for_export(self, charts: Dict) -> Dict[str, bytes]:
        """
        Process charts for export to different formats.
        
        Args:
            charts: Dictionary containing chart data
            
        Returns:
            Dictionary of chart images as bytes
        """
        chart_images = {}
        
        try:
            # Process each chart type
            if charts.get('themes_bar'):
                chart_images['themes_bar'] = self._convert_chart_to_image(charts['themes_bar'])
                
            if charts.get('classification_gauge'):
                chart_images['classification_gauge'] = self._convert_chart_to_image(charts['classification_gauge'])
                
            if charts.get('themes_pie'):
                chart_images['themes_pie'] = self._convert_chart_to_image(charts['themes_pie'])
                
            if charts.get('ethics_radar'):
                chart_images['ethics_radar'] = self._convert_chart_to_image(charts['ethics_radar'])
                
            logger.info("Processed %d charts for export", len(chart_images))
        except Exception as e:
            logger.warning("Error processing charts for export: %s", str(e))
            
        return chart_images
    
    def export_to_pdf(self, data: Dict[str, Any]) -> bytes:
        """
        Export recommendations and analysis results to PDF.
        
        Args:
            data: Dictionary containing analysis and recommendations data
            
        Returns:
            PDF file as bytes
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=72)
                               
        # Process charts if available
        chart_images = {}
        if data.get('charts'):
            chart_images = self._process_charts_for_export(data['charts'])
        
        # Create styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor(self.title_colour),
            spaceAfter=12
        )
        heading1_style = ParagraphStyle(
            'Heading1',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor(self.heading_colour),
            spaceBefore=12,
            spaceAfter=6
        )
        heading2_style = ParagraphStyle(
            'Heading2',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor(self.heading_colour),
            spaceBefore=10,
            spaceAfter=6
        )
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor(self.text_colour),
            spaceBefore=6,
            spaceAfter=6
        )
        
        # Build document content
        content = []
        
        # Title
        content.append(Paragraph("PolicyCraft Analysis Report", title_style))
        content.append(Spacer(1, 12))
        
        # Narrative (if available)
        if data.get('narrative') and data['narrative'].get('html'):
            narr_html = data['narrative']['html']
            content.append(Paragraph("Recommendation", heading1_style))
            # Very light HTML to paragraphs conversion
            try:
                # Split by block tags to keep some structure
                blocks = re.split(r"</?(h[1-4]|p|ul|ol|li|hr)[^>]*>", narr_html, flags=re.IGNORECASE)
                # Fallback: if split failed, just strip tags
                if len(blocks) <= 1:
                    text = re.sub(r"<[^>]+>", "", narr_html)
                    content.append(Paragraph(self._clean_text(text), normal_style))
                else:
                    # Remove tags and add as paragraphs; simple handling for list items
                    # Replace <li> with bullets by prefixing •
                    cleaned = re.sub(r"<li[^>]*>", "• ", narr_html, flags=re.IGNORECASE)
                    cleaned = re.sub(r"</li>", "\n", cleaned, flags=re.IGNORECASE)
                    cleaned = re.sub(r"<br\s*/?>", "\n", cleaned, flags=re.IGNORECASE)
                    # Strip remaining tags
                    cleaned = re.sub(r"<[^>]+>", "", cleaned)
                    for line in [l.strip() for l in cleaned.splitlines() if l.strip()]:
                        content.append(Paragraph(self._clean_text(line), normal_style))
            except Exception as _e:
                logger.warning(f"Failed to parse narrative HTML for PDF: {_e}")
            content.append(Spacer(1, 12))

        # Document information
        if data and data.get('analysis'):
            analysis = data['analysis']
            content.append(Paragraph("Document Information", heading1_style))
            
            # Create document info table
            doc_info = [
                ["Document", analysis.get('filename', 'Unknown')],
                ["Classification", analysis.get('classification', 'Unknown')],
                ["Generated Date", self._format_date(data.get('generated_date', 'Unknown'))],
                ["Total Recommendations", str(len(data.get('recommendations', [])))],
            ]
            
            doc_table = Table(doc_info, colWidths=[100, 350])
            doc_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#f8f9fa")),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor(self.heading_colour)),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e1e8ed")),
            ]))
            
            content.append(doc_table)
            content.append(Spacer(1, 12))
            
            # Analysis Results
            if analysis.get('themes'):
                content.append(Paragraph("Analysis Results", heading1_style))
                content.append(Paragraph("Key Themes", heading2_style))
                
                # Create themes table
                themes_data = [["Theme", "Score", "Confidence"]]
                for theme in analysis.get('themes', [])[:8]:  # Top 8 themes
                    themes_data.append([
                        theme.get('name', 'Unknown'),
                        str(theme.get('score', 0)),
                        f"{theme.get('confidence', 0)}%"
                    ])
                
                themes_table = Table(themes_data, colWidths=[250, 100, 100])
                themes_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f8f9fa")),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor(self.heading_colour)),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e1e8ed")),
                ]))
                
                content.append(themes_table)
                content.append(Spacer(1, 12))
                
                # Add charts if available
                if chart_images and any(v for v in chart_images.values()):
                    content.append(Paragraph("Visualisations", heading2_style))
                    
                    # Themes bar chart
                    if chart_images.get('themes_bar'):
                        content.append(Paragraph("Theme Distribution", ParagraphStyle(
                            'ChartTitle',
                            parent=normal_style,
                            fontName='Helvetica-Bold'
                        )))
                        img = Image(io.BytesIO(chart_images['themes_bar']))
                        img.drawHeight = 3 * inch
                        img.drawWidth = 5 * inch
                        content.append(img)
                        content.append(Spacer(1, 6))
                    
                    # Classification gauge
                    if chart_images.get('classification_gauge'):
                        content.append(Paragraph("Policy Classification", ParagraphStyle(
                            'ChartTitle',
                            parent=normal_style,
                            fontName='Helvetica-Bold'
                        )))
                        img = Image(io.BytesIO(chart_images['classification_gauge']))
                        img.drawHeight = 3 * inch
                        img.drawWidth = 5 * inch
                        content.append(img)
                        content.append(Spacer(1, 6))
                    
                    # Themes pie chart
                    if chart_images.get('themes_pie'):
                        content.append(Paragraph("Theme Proportions", ParagraphStyle(
                            'ChartTitle',
                            parent=normal_style,
                            fontName='Helvetica-Bold'
                        )))
                        img = Image(io.BytesIO(chart_images['themes_pie']))
                        img.drawHeight = 3 * inch
                        img.drawWidth = 5 * inch
                        content.append(img)
                        content.append(Spacer(1, 6))

                else:
                    # Explicit notice when charts cannot be embedded
                    content.append(Paragraph("Visualisations", heading2_style))
                    content.append(Paragraph(
                        self._clean_text(
                            "Charts could not be embedded in this PDF. To enable chart export, install the 'kaleido' package (pip install -U kaleido) and regenerate the report."
                        ),
                        normal_style,
                    ))
                    
                    # Ethics radar chart
                    if chart_images.get('ethics_radar'):
                        content.append(Paragraph("Ethical Dimensions Coverage", ParagraphStyle(
                            'ChartTitle',
                            parent=normal_style,
                            fontName='Helvetica-Bold'
                        )))
                        img = Image(io.BytesIO(chart_images['ethics_radar']))
                        img.drawHeight = 3 * inch
                        img.drawWidth = 5 * inch
                        content.append(img)
                        content.append(Spacer(1, 12))
            
            # Recommendations
            content.append(Paragraph("Strategic Recommendations", heading1_style))
            
            if data.get('recommendations'):
                for i, rec in enumerate(data.get('recommendations', []), 1):
                    # Recommendation title
                    title = rec.get('title', f'Recommendation {i}')
                    content.append(Paragraph(f"{i}. {title}", heading2_style))
                    
                    # Description
                    description = self._clean_text(rec.get('description', ''))
                    content.append(Paragraph(description, normal_style))
                    content.append(Spacer(1, 6))
                    
                    # Implementation steps
                    if rec.get('implementation_steps'):
                        content.append(Paragraph("Implementation Steps:", ParagraphStyle(
                            'Steps',
                            parent=normal_style,
                            fontName='Helvetica-Bold'
                        )))
                        
                        steps = []
                        for step in rec.get('implementation_steps', []):
                            steps.append(ListItem(Paragraph(self._clean_text(step), normal_style)))
                        
                        content.append(ListFlowable(
                            steps,
                            bulletType='1',
                            start=1,
                            bulletFontSize=10,
                            leftIndent=20,
                            bulletFormat='%s.'
                        ))
                        content.append(Spacer(1, 6))
                    
                    # Sources
                    if rec.get('sources'):
                        content.append(Paragraph("Sources:", ParagraphStyle(
                            'Sources',
                            parent=normal_style,
                            fontName='Helvetica-Bold'
                        )))
                        
                        sources = []
                        for source in rec.get('sources', []):
                            sources.append(ListItem(Paragraph(self._clean_text(source), ParagraphStyle(
                                'SourceItem',
                                parent=normal_style,
                                fontName='Helvetica-Oblique'
                            ))))
                    elif rec.get('source'):
                        content.append(Paragraph("Sources:", ParagraphStyle(
                            'Sources',
                            parent=normal_style,
                            fontName='Helvetica-Bold'
                        )))
                        content.append(ListFlowable(
                            [ListItem(Paragraph(self._clean_text(rec.get('source')), ParagraphStyle('SourceItem', parent=normal_style, fontName='Helvetica-Oblique')))],
                            bulletType='bullet',
                            leftIndent=20
                        ))
                        
                        content.append(ListFlowable(
                            sources,
                            bulletType='bullet',
                            leftIndent=20
                        ))
                        content.append(Spacer(1, 6))
                    
                    # Timeframe
                    if rec.get('timeframe'):
                        content.append(Paragraph(f"Timeframe: {rec.get('timeframe')}", ParagraphStyle(
                            'Timeframe',
                            parent=normal_style,
                            fontName='Helvetica-Bold',
                            textColor=colors.HexColor(self.accent_colour)
                        )))
                    
                    content.append(Spacer(1, 12))
            else:
                content.append(Paragraph("No specific recommendations are available for this policy document.", normal_style))
            
            # Add missing sections from recommendations template
            
            # Methodology & Confidence
            content.append(Paragraph("Methodology & Confidence", heading1_style))
            content.append(Paragraph("Methodology", heading2_style))
            methodology_text = data.get('methodology', 'PolicyCraft local analysis pipeline (text extraction → classification → theme detection → rules‑based + ML heuristics).')
            content.append(Paragraph(self._clean_text(methodology_text), normal_style))
            content.append(Paragraph("Note: This report relies solely on locally available data (no external benchmarks).", ParagraphStyle(
                'Note',
                parent=normal_style,
                fontName='Helvetica-Oblique'
            )))
            
            content.append(Paragraph("Confidence", heading2_style))
            confidence_pct = data.get('analysis', {}).get('confidence_pct', 0)
            content.append(Paragraph(f"Overall confidence: {confidence_pct:.0f}%", normal_style))
            content.append(Paragraph("• Number of independent evidence snippets in the document", normal_style))
            content.append(Paragraph("• Length and quality of extracted text", normal_style))
            content.append(Paragraph("• Presence of formal policy structure", normal_style))
            content.append(Spacer(1, 12))
            
            # Benchmarks
            content.append(Paragraph("Benchmarks", heading1_style))
            cf = data.get('analysis', {}).get('confidence_factors', {})
            
            benchmarks_data = [["Metric", "Value"]]
            benchmarks_data.append(["Unique sources", str(cf.get('unique_sources', '—'))])
            benchmarks_data.append(["Evidence diversity", f"{cf.get('evidence_diversity', 0)}%"])
            benchmarks_data.append(["Text quality score", f"{cf.get('text_quality', 0)}%"])
            benchmarks_data.append(["Coverage of target length", f"{cf.get('text_quality', 0)}%"])
            
            benchmarks_table = Table(benchmarks_data, colWidths=[200, 200])
            benchmarks_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f8f9fa")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor(self.heading_colour)),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e1e8ed")),
            ]))
            content.append(benchmarks_table)
            content.append(Spacer(1, 12))
            
            # Stakeholder Perspectives
            content.append(Paragraph("Stakeholder Perspectives", heading1_style))
            
            content.append(Paragraph("Students", heading2_style))
            content.append(Paragraph("• Transparency of AI usage rules and appeal processes", normal_style))
            content.append(Paragraph("• Clear guides on permitted vs non‑permitted use", normal_style))
            
            content.append(Paragraph("Faculty", heading2_style))
            content.append(Paragraph("• Training on AI detection and assessment methods", normal_style))
            content.append(Paragraph("• Clear protocols for handling suspected violations", normal_style))
            
            content.append(Paragraph("Administration", heading2_style))
            content.append(Paragraph("• Streamlined investigation procedures", normal_style))
            content.append(Paragraph("• Regular review and updates of AI usage policies", normal_style))
            content.append(Spacer(1, 12))
            
            # Pilot & Evaluation Plan
            content.append(Paragraph("Pilot & Evaluation Plan", heading1_style))
            content.append(Paragraph("• 30 days: pilot in selected courses; metrics: compliance, incident rate, satisfaction", normal_style))
            content.append(Paragraph("• 60 days: scale‑up, refine guidelines, follow‑up training", normal_style))
            content.append(Paragraph("• 90 days: evaluate results; go/no‑go decision for full roll‑out", normal_style))
            content.append(Paragraph("90‑day targets are treated as pilot goals – full institutionalisation after results review.", ParagraphStyle(
                'Note',
                parent=normal_style,
                fontName='Helvetica-Oblique'
            )))
            content.append(Spacer(1, 12))
            
            # Impact–Urgency–Feasibility Matrix
            content.append(Paragraph("Impact–Urgency–Feasibility", heading1_style))
            
            # Create IUF table for recommendations
            if data.get('recommendations'):
                iuf_data = [["Recommendation", "Impact", "Urgency", "Feasibility"]]
                for i, rec in enumerate(data.get('recommendations', [])[:5], 1):  # Top 5
                    title = rec.get('title', f'Recommendation {i}')
                    # Use priority to determine values, or use defaults
                    priority = rec.get('priority', 'medium').lower()
                    if priority == 'high':
                        impact, urgency, feasibility = "High", "High", "Medium"
                    elif priority == 'low':
                        impact, urgency, feasibility = "Medium", "Low", "High"
                    else:
                        impact, urgency, feasibility = "Medium", "Medium", "Medium"
                    
                    iuf_data.append([title[:40] + "..." if len(title) > 40 else title, impact, urgency, feasibility])
                
                iuf_table = Table(iuf_data, colWidths=[200, 80, 80, 80])
                iuf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f8f9fa")),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor(self.heading_colour)),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e1e8ed")),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                content.append(iuf_table)
            else:
                content.append(Paragraph("No recommendations available for impact-urgency-feasibility analysis.", normal_style))
            
            content.append(Spacer(1, 12))
            
            # Footer
            content.append(Spacer(1, 24))
            content.append(Paragraph(f"PolicyCraft Analysis Report - {self._format_date(data.get('generated_date', datetime.now().isoformat()))}", ParagraphStyle(
                'Footer',
                parent=normal_style,
                fontSize=8,
                textColor=colors.HexColor("#7f8c8d"),
                alignment=1  # Centre alignment
            )))
        else:
            content.append(Paragraph("No data available for export.", normal_style))
        
        # Build PDF
        doc.build(content)
        pdf_data = buffer.getvalue()
        buffer.close()
        
        return pdf_data
    
    def export_to_word(self, data: Dict[str, Any]) -> bytes:
        """
        Export recommendations and analysis results to Word document.
        
        Args:
            data: Dictionary containing analysis and recommendations data
            
        Returns:
            Word document as bytes
        """
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "PolicyCraft Analysis Report"
        doc.core_properties.author = "PolicyCraft"
        doc.core_properties.created = datetime.now()
        
        # Process charts if available
        chart_images = {}
        if data.get('charts'):
            chart_images = self._process_charts_for_export(data['charts'])
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)
        
        # Title
        title = doc.add_heading("PolicyCraft Analysis Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style.font.color.rgb = RGBColor.from_string(self.title_colour.replace('#', ''))
        
        # Narrative (if available)
        if data.get('narrative') and data['narrative'].get('html'):
            doc.add_heading("Recommendation", level=1)
            narr_html = data['narrative']['html']
            try:
                # Replace list items and breaks to preserve structure
                cleaned = re.sub(r"<li[^>]*>", "- ", narr_html, flags=re.IGNORECASE)
                cleaned = re.sub(r"</li>", "\n", cleaned, flags=re.IGNORECASE)
                cleaned = re.sub(r"<br\s*/?>", "\n", cleaned, flags=re.IGNORECASE)
                cleaned = re.sub(r"<[^>]+>", "", cleaned)
                for line in [l.strip() for l in cleaned.splitlines() if l.strip()]:
                    doc.add_paragraph(self._clean_text(line))
            except Exception as _e:
                logger.warning(f"Failed to parse narrative HTML for Word: {_e}")
            doc.add_paragraph()
        
        # Document information
        if data and data.get('analysis'):
            analysis = data['analysis']
            doc.add_heading("Document Information", level=1)
            
            # Create document info table
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            
            # Add data to table
            rows = table.rows
            rows[0].cells[0].text = "Document"
            rows[0].cells[1].text = analysis.get('filename', 'Unknown')
            rows[1].cells[0].text = "Classification"
            rows[1].cells[1].text = analysis.get('classification', 'Unknown')
            rows[2].cells[0].text = "Generated Date"
            rows[2].cells[1].text = self._format_date(data.get('generated_date', 'Unknown'))
            rows[3].cells[0].text = "Total Recommendations"
            rows[3].cells[1].text = str(len(data.get('recommendations', [])))
            
            # Format table
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            
            doc.add_paragraph()
            
            # Add visualisations section if charts are available
            if chart_images:
                doc.add_heading('Visualisations', level=2)
                
                # Themes bar chart
                if chart_images.get('themes_bar'):
                    doc.add_paragraph('Theme Distribution', style='Heading 3')
                    doc.add_picture(io.BytesIO(chart_images['themes_bar']), width=Cm(15))
                    doc.add_paragraph()
                
                # Classification gauge
                if chart_images.get('classification_gauge'):
                    doc.add_paragraph('Policy Classification', style='Heading 3')
                    doc.add_picture(io.BytesIO(chart_images['classification_gauge']), width=Cm(15))
                    doc.add_paragraph()
                
                # Themes pie chart
                if chart_images.get('themes_pie'):
                    doc.add_paragraph('Theme Proportions', style='Heading 3')
                    doc.add_picture(io.BytesIO(chart_images['themes_pie']), width=Cm(15))
                    doc.add_paragraph()
                
                # Ethics radar chart
                if chart_images.get('ethics_radar'):
                    doc.add_paragraph('Ethical Dimensions Coverage', style='Heading 3')
                    doc.add_picture(io.BytesIO(chart_images['ethics_radar']), width=Cm(15))
                    doc.add_paragraph()
                
                doc.add_paragraph()
            
            # Recommendations
            doc.add_heading("Strategic Recommendations", level=1)
            
            if data.get('recommendations'):
                for i, rec in enumerate(data.get('recommendations', []), 1):
                    # Recommendation title
                    title = rec.get('title', f'Recommendation {i}')
                    doc.add_heading(f"{i}. {title}", level=2)
                    
                    # Description
                    doc.add_paragraph(self._clean_text(rec.get('description', '')))
                    
                    # Implementation steps
                    if rec.get('implementation_steps'):
                        p = doc.add_paragraph()
                        p.add_run("Implementation Steps:").bold = True
                        
                        for j, step in enumerate(rec.get('implementation_steps', []), 1):
                            doc.add_paragraph(self._clean_text(f"{j}. {step}"), style='List Number')
                    
                    # Sources
                    if rec.get('sources'):
                        p = doc.add_paragraph()
                        p.add_run("Sources:").bold = True
                        
                        for source in rec.get('sources', []):
                            p = doc.add_paragraph(self._clean_text(source), style='List Bullet')
                            p.style.font.italic = True
                    elif rec.get('source'):
                        p = doc.add_paragraph()
                        p.add_run("Sources:").bold = True
                        q = doc.add_paragraph(self._clean_text(rec.get('source')), style='List Bullet')
                        q.style.font.italic = True
                    
                    # Timeframe
                    if rec.get('timeframe'):
                        p = doc.add_paragraph()
                        run = p.add_run(f"Timeframe: {rec.get('timeframe')}")
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(self.accent_colour.replace('#', ''))
                    
                    doc.add_paragraph()
            else:
                doc.add_paragraph("No specific recommendations are available for this policy document.")
            
            # Add missing sections from recommendations template
            
            # Methodology & Confidence
            doc.add_heading("Methodology & Confidence", level=1)
            doc.add_heading("Methodology", level=2)
            methodology_text = data.get('methodology', 'PolicyCraft local analysis pipeline (text extraction → classification → theme detection → rules‑based + ML heuristics).')
            doc.add_paragraph(methodology_text)
            note_para = doc.add_paragraph("Note: This report relies solely on locally available data (no external benchmarks).")
            note_para.italic = True
            
            doc.add_heading("Confidence", level=2)
            confidence_pct = data.get('analysis', {}).get('confidence_pct', 0)
            doc.add_paragraph(f"Overall confidence: {confidence_pct:.0f}%")
            doc.add_paragraph("• Number of independent evidence snippets in the document")
            doc.add_paragraph("• Length and quality of extracted text")
            doc.add_paragraph("• Presence of formal policy structure")
            
            # Benchmarks
            doc.add_heading("Benchmarks", level=1)
            cf = data.get('analysis', {}).get('confidence_factors', {})
            
            # Create benchmarks table
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Light Grid Accent 1'
            table.cell(0, 0).text = "Metric"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "Unique sources"
            table.cell(1, 1).text = str(cf.get('unique_sources', '—'))
            table.cell(2, 0).text = "Evidence diversity"
            table.cell(2, 1).text = f"{cf.get('evidence_diversity', 0)}%"
            table.cell(3, 0).text = "Text quality score"
            table.cell(3, 1).text = f"{cf.get('text_quality', 0)}%"
            table.cell(4, 0).text = "Coverage of target length"
            table.cell(4, 1).text = f"{cf.get('text_quality', 0)}%"
            
            # Stakeholder Perspectives
            doc.add_heading("Stakeholder Perspectives", level=1)
            
            doc.add_heading("Students", level=2)
            doc.add_paragraph("• Transparency of AI usage rules and appeal processes")
            doc.add_paragraph("• Clear guides on permitted vs non‑permitted use")
            
            doc.add_heading("Faculty", level=2)
            doc.add_paragraph("• Training on AI detection and assessment methods")
            doc.add_paragraph("• Clear protocols for handling suspected violations")
            
            doc.add_heading("Administration", level=2)
            doc.add_paragraph("• Streamlined investigation procedures")
            doc.add_paragraph("• Regular review and updates of AI usage policies")
            
            # Pilot & Evaluation Plan
            doc.add_heading("Pilot & Evaluation Plan", level=1)
            doc.add_paragraph("• 30 days: pilot in selected courses; metrics: compliance, incident rate, satisfaction")
            doc.add_paragraph("• 60 days: scale‑up, refine guidelines, follow‑up training")
            doc.add_paragraph("• 90 days: evaluate results; go/no‑go decision for full roll‑out")
            note_para = doc.add_paragraph("90‑day targets are treated as pilot goals – full institutionalisation after results review.")
            note_para.italic = True
            
            # Impact–Urgency–Feasibility Matrix
            doc.add_heading("Impact–Urgency–Feasibility", level=1)
            
            # Create IUF table for recommendations
            if data.get('recommendations'):
                num_recs = min(len(data.get('recommendations', [])), 5)  # Top 5
                iuf_table = doc.add_table(rows=num_recs + 1, cols=4)
                iuf_table.style = 'Light Grid Accent 1'
                
                # Header row
                iuf_table.cell(0, 0).text = "Recommendation"
                iuf_table.cell(0, 1).text = "Impact"
                iuf_table.cell(0, 2).text = "Urgency"
                iuf_table.cell(0, 3).text = "Feasibility"
                
                # Data rows
                for i, rec in enumerate(data.get('recommendations', [])[:5], 1):
                    title = rec.get('title', f'Recommendation {i}')
                    # Use priority to determine values, or use defaults
                    priority = rec.get('priority', 'medium').lower()
                    if priority == 'high':
                        impact, urgency, feasibility = "High", "High", "Medium"
                    elif priority == 'low':
                        impact, urgency, feasibility = "Medium", "Low", "High"
                    else:
                        impact, urgency, feasibility = "Medium", "Medium", "Medium"
                    
                    iuf_table.cell(i, 0).text = title[:60] + "..." if len(title) > 60 else title
                    iuf_table.cell(i, 1).text = impact
                    iuf_table.cell(i, 2).text = urgency
                    iuf_table.cell(i, 3).text = feasibility
            else:
                doc.add_paragraph("No recommendations available for impact-urgency-feasibility analysis.")
            
            # Footer
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"PolicyCraft Analysis Report - {self._format_date(data.get('generated_date', datetime.now().isoformat()))}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph("No data available for export.")
        
        # Save document to memory
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_data = buffer.getvalue()
        buffer.close()
        
        return docx_data
    
    def export_to_excel(self, data: Dict[str, Any]) -> bytes:
        """
        Export recommendations and analysis results to Excel spreadsheet.
        
        Args:
            data: Dictionary containing analysis and recommendations data
            
        Returns:
            Excel spreadsheet as bytes
        """
        buffer = io.BytesIO()
        workbook = xlsxwriter.Workbook(buffer)
        
        # Create formats
        title_format = workbook.add_format({
            'bold': True,
            'font_size': 16,
            'font_color': self.title_colour,
            'align': 'centre',
            'valign': 'vcenter'
        })
        
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#f8f9fa',
            'border': 1,
            'border_color': '#e1e8ed',
            'font_color': self.heading_colour,
            'align': 'centre',
            'valign': 'vcenter'
        })
        
        cell_format = workbook.add_format({
            'border': 1,
            'border_color': '#e1e8ed',
            'font_color': self.text_colour,
            'align': 'left',
            'valign': 'vcenter',
            'text_wrap': True
        })
        
        # Document Information sheet
        info_sheet = workbook.add_worksheet('Document Info')
        info_sheet.set_column('A:A', 20)
        info_sheet.set_column('B:B', 50)
        
        # Add title
        info_sheet.merge_range('A1:B1', 'PolicyCraft Analysis Report', title_format)
        
        if data and data.get('analysis'):
            analysis = data['analysis']
            
            # Add document info
            info_sheet.write('A3', 'Document', header_format)
            info_sheet.write('B3', analysis.get('filename', 'Unknown'), cell_format)
            
            info_sheet.write('A4', 'Classification', header_format)
            info_sheet.write('B4', analysis.get('classification', 'Unknown'), cell_format)
            
            info_sheet.write('A5', 'Generated Date', header_format)
            info_sheet.write('B5', self._format_date(data.get('generated_date', 'Unknown')), cell_format)
            
            info_sheet.write('A6', 'Total Recommendations', header_format)
            info_sheet.write('B6', len(data.get('recommendations', [])), cell_format)
            
            # Themes sheet
            if analysis.get('themes'):
                themes_sheet = workbook.add_worksheet('Themes')
                themes_sheet.set_column('A:A', 40)
                themes_sheet.set_column('B:B', 15)
                themes_sheet.set_column('C:C', 15)
                
                # Add title
                themes_sheet.merge_range('A1:C1', 'Key Themes Analysis', title_format)
                
                # Add header row
                themes_sheet.write('A3', 'Theme', header_format)
                themes_sheet.write('B3', 'Score', header_format)
                themes_sheet.write('C3', 'Confidence', header_format)
                
                # Add theme data
                row = 3
                for theme in analysis.get('themes', []):
                    themes_sheet.write(row, 0, theme.get('name', 'Unknown'), cell_format)
                    themes_sheet.write(row, 1, theme.get('score', 0), cell_format)
                    themes_sheet.write(row, 2, f"{theme.get('confidence', 0)}%", cell_format)
                    row += 1
            
            # Recommendations sheet
            recs_sheet = workbook.add_worksheet('Recommendations')
            recs_sheet.set_column('A:A', 5)
            recs_sheet.set_column('B:B', 30)
            recs_sheet.set_column('C:C', 50)
            recs_sheet.set_column('D:D', 50)
            recs_sheet.set_column('E:E', 50)
            recs_sheet.set_column('F:F', 15)
            
            # Add title
            recs_sheet.merge_range('A1:F1', 'Strategic Recommendations', title_format)
            
            # Add header row
            recs_sheet.write('A3', '#', header_format)
            recs_sheet.write('B3', 'Title', header_format)
            recs_sheet.write('C3', 'Description', header_format)
            recs_sheet.write('D3', 'Implementation Steps', header_format)
            recs_sheet.write('E3', 'Sources', header_format)
            recs_sheet.write('F3', 'Timeframe', header_format)
            
            # Add recommendation data
            row = 3
            for i, rec in enumerate(data.get('recommendations', []), 1):
                recs_sheet.write(row, 0, i, cell_format)
                recs_sheet.write(row, 1, rec.get('title', f'Recommendation {i}'), cell_format)
                recs_sheet.write(row, 2, self._clean_text(rec.get('description', '')), cell_format)
                
                # Implementation steps
                steps_text = ""
                for j, step in enumerate(rec.get('implementation_steps', []), 1):
                    steps_text += self._clean_text(f"{j}. {step}") + "\n"
                recs_sheet.write(row, 3, steps_text, cell_format)
                
                # Sources
                sources_text = ""
                if rec.get('sources'):
                    for source in rec.get('sources', []):
                        sources_text += "• " + self._clean_text(source) + "\n"
                elif rec.get('source'):
                    sources_text = "• " + self._clean_text(rec.get('source'))
                recs_sheet.write(row, 4, sources_text, cell_format)
                
                # Timeframe
                recs_sheet.write(row, 5, rec.get('timeframe', ''), cell_format)
                
                row += 1
        else:
            info_sheet.write('A3', 'No data available for export.', cell_format)
        
        # Add missing sections as new worksheets
        
        # Methodology & Confidence sheet
        methodology_sheet = workbook.add_worksheet('Methodology & Confidence')
        methodology_sheet.set_column('A:A', 30)
        methodology_sheet.set_column('B:B', 60)
        
        methodology_sheet.write('A1', 'Methodology & Confidence', title_format)
        methodology_sheet.write('A3', 'Methodology', header_format)
        methodology_text = data.get('methodology', 'PolicyCraft local analysis pipeline (text extraction → classification → theme detection → rules‑based + ML heuristics).')
        methodology_sheet.write('B3', methodology_text, cell_format)
        methodology_sheet.write('B4', 'Note: This report relies solely on locally available data (no external benchmarks).', cell_format)
        
        methodology_sheet.write('A6', 'Confidence', header_format)
        confidence_pct = data.get('analysis', {}).get('confidence_pct', 0)
        methodology_sheet.write('B6', f'Overall confidence: {confidence_pct:.0f}%', cell_format)
        methodology_sheet.write('B7', '• Number of independent evidence snippets in the document', cell_format)
        methodology_sheet.write('B8', '• Length and quality of extracted text', cell_format)
        methodology_sheet.write('B9', '• Presence of formal policy structure', cell_format)
        
        # Benchmarks sheet
        benchmarks_sheet = workbook.add_worksheet('Benchmarks')
        benchmarks_sheet.set_column('A:A', 30)
        benchmarks_sheet.set_column('B:B', 20)
        
        benchmarks_sheet.write('A1', 'Benchmarks', title_format)
        benchmarks_sheet.write('A3', 'Metric', header_format)
        benchmarks_sheet.write('B3', 'Value', header_format)
        
        cf = data.get('analysis', {}).get('confidence_factors', {})
        benchmarks_sheet.write('A4', 'Unique sources', cell_format)
        benchmarks_sheet.write('B4', str(cf.get('unique_sources', '—')), cell_format)
        benchmarks_sheet.write('A5', 'Evidence diversity', cell_format)
        benchmarks_sheet.write('B5', f"{cf.get('evidence_diversity', 0)}%", cell_format)
        benchmarks_sheet.write('A6', 'Text quality score', cell_format)
        benchmarks_sheet.write('B6', f"{cf.get('text_quality', 0)}%", cell_format)
        benchmarks_sheet.write('A7', 'Coverage of target length', cell_format)
        benchmarks_sheet.write('B7', f"{cf.get('text_quality', 0)}%", cell_format)
        
        # Stakeholder Perspectives sheet
        stakeholders_sheet = workbook.add_worksheet('Stakeholder Perspectives')
        stakeholders_sheet.set_column('A:A', 20)
        stakeholders_sheet.set_column('B:B', 80)
        
        stakeholders_sheet.write('A1', 'Stakeholder Perspectives', title_format)
        stakeholders_sheet.write('A3', 'Students', header_format)
        stakeholders_sheet.write('B3', '• Transparency of AI usage rules and appeal processes', cell_format)
        stakeholders_sheet.write('B4', '• Clear guides on permitted vs non‑permitted use', cell_format)
        
        stakeholders_sheet.write('A6', 'Faculty', header_format)
        stakeholders_sheet.write('B6', '• Training on AI detection and assessment methods', cell_format)
        stakeholders_sheet.write('B7', '• Clear protocols for handling suspected violations', cell_format)
        
        stakeholders_sheet.write('A9', 'Administration', header_format)
        stakeholders_sheet.write('B9', '• Streamlined investigation procedures', cell_format)
        stakeholders_sheet.write('B10', '• Regular review and updates of AI usage policies', cell_format)
        
        # Pilot & Evaluation Plan sheet
        pilot_sheet = workbook.add_worksheet('Pilot & Evaluation Plan')
        pilot_sheet.set_column('A:A', 100)
        
        pilot_sheet.write('A1', 'Pilot & Evaluation Plan', title_format)
        pilot_sheet.write('A3', '• 30 days: pilot in selected courses; metrics: compliance, incident rate, satisfaction', cell_format)
        pilot_sheet.write('A4', '• 60 days: scale‑up, refine guidelines, follow‑up training', cell_format)
        pilot_sheet.write('A5', '• 90 days: evaluate results; go/no‑go decision for full roll‑out', cell_format)
        pilot_sheet.write('A7', '90‑day targets are treated as pilot goals – full institutionalisation after results review.', cell_format)
        
        # Impact–Urgency–Feasibility Matrix sheet
        iuf_sheet = workbook.add_worksheet('Impact-Urgency-Feasibility')
        iuf_sheet.set_column('A:A', 40)
        iuf_sheet.set_column('B:D', 15)
        
        iuf_sheet.write('A1', 'Impact–Urgency–Feasibility Matrix', title_format)
        iuf_sheet.write('A3', 'Recommendation', header_format)
        iuf_sheet.write('B3', 'Impact', header_format)
        iuf_sheet.write('C3', 'Urgency', header_format)
        iuf_sheet.write('D3', 'Feasibility', header_format)
        
        if data.get('recommendations'):
            row = 4
            for i, rec in enumerate(data.get('recommendations', [])[:10], 1):  # Top 10 for Excel
                title = rec.get('title', f'Recommendation {i}')
                # Use priority to determine values, or use defaults
                priority = rec.get('priority', 'medium').lower()
                if priority == 'high':
                    impact, urgency, feasibility = "High", "High", "Medium"
                elif priority == 'low':
                    impact, urgency, feasibility = "Medium", "Low", "High"
                else:
                    impact, urgency, feasibility = "Medium", "Medium", "Medium"
                
                iuf_sheet.write(row, 0, title, cell_format)
                iuf_sheet.write(row, 1, impact, cell_format)
                iuf_sheet.write(row, 2, urgency, cell_format)
                iuf_sheet.write(row, 3, feasibility, cell_format)
                row += 1
        else:
            iuf_sheet.write('A4', 'No recommendations available for impact-urgency-feasibility analysis.', cell_format)
        
        # Charts Information sheet (since Excel can't easily embed Plotly charts)
        charts_sheet = workbook.add_worksheet('Charts Information')
        charts_sheet.set_column('A:A', 100)
        
        charts_sheet.write('A1', 'Charts Information', title_format)
        if data.get('charts'):
            charts_sheet.write('A3', 'Available charts in the PolicyCraft analysis (not embedded in Excel):', cell_format)
            row = 4
            if data['charts'].get('themes_bar'):
                charts_sheet.write(row, 0, '• Theme Distribution Bar Chart - shows relative importance of policy themes', cell_format)
                row += 1
            if data['charts'].get('classification_gauge'):
                charts_sheet.write(row, 0, '• Policy Classification Gauge - indicates policy type and confidence level', cell_format)
                row += 1
            if data['charts'].get('themes_pie'):
                charts_sheet.write(row, 0, '• Theme Proportions Pie Chart - displays theme distribution as percentages', cell_format)
                row += 1
            if data['charts'].get('ethics_radar'):
                charts_sheet.write(row, 0, '• Ethical Dimensions Radar Chart - shows coverage of ethical considerations', cell_format)
                row += 1
            
            row += 1
            charts_sheet.write(row, 0, 'Note: To view interactive charts, please use the PDF or Word export, or access the web interface.', cell_format)
        else:
            charts_sheet.write('A3', 'No charts are available for this analysis.', cell_format)
        
        workbook.close()
        excel_data = buffer.getvalue()
        buffer.close()
        
        return excel_data
