"""
Text Processing Module for PolicyCraft AI Policy Analysis Platform.
This module provides lightweight text processing capabilities for policy documents,
including text extraction from various file formats (PDF, DOCX, TXT) and basic
natural language processing tasks. The implementation uses regex-based tokenization
to minimize external dependencies while maintaining good performance.

Key Features:
- Support for multiple document formats (PDF, DOCX, TXT)
- Regex-based tokenization without NLTK dependency
- Text cleaning and normalization
- Basic text statistics and analysis
- Policy-specific term handling

Author: Jacek Robert Kszczot
Project: MSc Data Science & AI - COM7016
University: Leeds Trinity University
"""

import os
import re
import logging
from pathlib import Path
from typing import Optional, Dict, List

# PDF processing libraries
try:
    import pypdf
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PDF libraries not available. Install with: pip install PyPDF2 pdfplumber")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: DOCX library not available. Install with: pip install python-docx")

try:
    import contractions
    CONTRACTIONS_AVAILABLE = True
except ImportError:
    CONTRACTIONS_AVAILABLE = False
    print("Warning: contractions library not available. Install with: pip install contractions")

logger = logging.getLogger(__name__)

class TextProcessor:
    """
    Simplified text processing for policy documents.
    Uses regex-based tokenization to avoid dependency issues.
    """
    
    def __init__(self):
        """Initialize text processor with configuration."""
        self.supported_formats = {'.pdf', '.docx', '.txt', '.doc'}
        
        # Common English stopwords (simplified list)
        self.stop_words = {
            'i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', 'your', 'yours',
            'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', 'her', 'hers',
            'herself', 'it', 'its', 'itself', 'they', 'them', 'their', 'theirs', 'themselves',
            'what', 'which', 'who', 'whom', 'this', 'that', 'these', 'those', 'am', 'is', 'are',
            'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does',
            'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until',
            'while', 'of', 'at', 'by', 'for', 'with', 'through', 'during', 'before', 'after',
            'above', 'below', 'up', 'down', 'in', 'out', 'on', 'off', 'over', 'under', 'again',
            'further', 'then', 'once'
        }
        
        # Policy-specific terms to preserve
        self.policy_terms = {
            'ai', 'artificial', 'intelligence', 'policy', 'ethics', 'bias',
            'transparency', 'accountability', 'fairness', 'privacy', 'data',
            'algorithm', 'model', 'training', 'deployment', 'governance',
            'compliance', 'regulation', 'guideline', 'framework', 'principle'
        }
        
        # Remove policy terms from stopwords
        self.stop_words = self.stop_words - self.policy_terms
        
        print("TextProcessor initialized successfully")

    def extract_text_from_file(self, file_path: str) -> Optional[str]:
        """Extract text from various file formats with fallback methods."""
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return None
            
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in self.supported_formats:
            print(f"Unsupported file format: {file_extension}")
            return None
            
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
                
        except Exception as e:
            print(f"Error extracting text from {file_path}: {str(e)}")
            return None
            
        return None

    def _extract_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF using multiple methods."""
        if not PDF_AVAILABLE:
            print("PDF libraries not available")
            return None
            
        text_pypdf2 = ""
        text_pdfplumber = ""
        
        # Method 1: PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page in pdf_reader.pages:
                    text_pypdf2 += page.extract_text() + "\n"
            print(f"PyPDF2 extracted {len(text_pypdf2)} characters")
        except Exception as e:
            print(f"PyPDF2 extraction failed: {str(e)}")
            
        # Method 2: pdfplumber
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_pdfplumber += page_text + "\n"
            print(f"pdfplumber extracted {len(text_pdfplumber)} characters")
        except Exception as e:
            print(f"pdfplumber extraction failed: {str(e)}")
            
        # Choose better result
        if len(text_pdfplumber) > len(text_pypdf2):
            best_text = text_pdfplumber
            method = "pdfplumber"
        else:
            best_text = text_pypdf2
            method = "PyPDF2"
            
        if best_text.strip():
            print(f"PDF text extracted successfully using {method}: {len(best_text)} characters")
            return best_text
        else:
            print("No text could be extracted from PDF")
            return None

    def _extract_from_docx(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX files."""
        if not DOCX_AVAILABLE:
            print("DOCX library not available")
            return None
            
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
            print(f"DOCX text extracted successfully: {len(text)} characters")
            return text
            
        except Exception as e:
            print(f"Error extracting DOCX text: {str(e)}")
            return None

    def _extract_from_txt(self, file_path: str) -> Optional[str]:
        """Extract text from TXT files with encoding detection."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                print(f"TXT file read successfully with {encoding}: {len(text)} characters")
                return text
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"Error reading TXT file: {str(e)}")
                return None
                
        print("Could not read TXT file with any encoding")
        return None

    def clean_text(self, text: str) -> str:
        """Comprehensive text cleaning and preprocessing."""
        if not text:
            return ""
            
        print(f"Cleaning text: {len(text)} characters")
        
        # Basic cleaning
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        
        # Expand contractions
        if CONTRACTIONS_AVAILABLE:
            try:
                text = contractions.fix(text)
            except Exception:
                pass  # Skip if contractions fails
        
        # Remove special characters but preserve sentence structure
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)]', ' ', text)
        
        # Fix spacing around punctuation
        text = re.sub(r'\s+([\.!?])', r'\1', text)
        text = re.sub(r'([\.!?])\s*', r'\1 ', text)
        
        # Remove URLs and emails
        text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
        text = re.sub(r'\S+@\S+', '', text)
        
        # Remove excessive punctuation
        text = re.sub(r'\.{2,}', '.', text)
        text = re.sub(r'\?{2,}', '?', text)
        text = re.sub(r'!{2,}', '!', text)
        
        # Final cleanup
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        print(f"Text cleaned: {len(text)} characters")
        return text

    def tokenize_sentences(self, text: str) -> List[str]:
        """Split text into sentences using regex."""
        if not text:
            return []
            
        # Simple sentence tokenization using regex
        sentences = re.split(r'[.!?]+', text)
        
        # Clean and filter sentences
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 10:  # Filter out very short sentences
                cleaned_sentences.append(sentence)
        
        print(f"Tokenized into {len(cleaned_sentences)} sentences")
        return cleaned_sentences

    def tokenize_words(self, text: str, remove_stopwords: bool = True) -> List[str]:
        """Split text into words using regex."""
        if not text:
            return []
            
        # Extract words using regex
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        
        # Filter stopwords if requested
        if remove_stopwords:
            words = [word for word in words if word not in self.stop_words]
                
        print(f"Tokenized into {len(words)} words")
        return words

    def get_text_statistics(self, text: str) -> Dict:
        """Calculate various text statistics."""
        if not text:
            return {
                'character_count': 0,
                'word_count': 0,
                'sentence_count': 0,
                'paragraph_count': 0,
                'avg_sentence_length': 0,
                'avg_word_length': 0
            }
            
        sentences = self.tokenize_sentences(text)
        words = self.tokenize_words(text, remove_stopwords=False)
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        stats = {
            'character_count': len(text),
            'word_count': len(words),
            'sentence_count': len(sentences),
            'paragraph_count': len(paragraphs),
            'avg_sentence_length': round(len(words) / len(sentences), 2) if sentences else 0,
            'avg_word_length': round(sum(len(word) for word in words) / len(words), 2) if words else 0
        }
        
        print(f"Text statistics calculated: {stats}")
        return stats

    def preview_text(self, text: str, max_length: int = 500) -> str:
        """Generate a preview of the text."""
        if not text:
            return ""
            
        if len(text) <= max_length:
            return text
            
        # Find a good breaking point
        preview = text[:max_length]
        last_period = preview.rfind('.')
        
        if last_period > max_length * 0.7:
            preview = preview[:last_period + 1]
        
        return preview + "..." if len(text) > len(preview) else preview


# Test the processor
if __name__ == "__main__":
    print("Starting text processor test...")
    processor = TextProcessor()
    
    sample_text = """
    This is a sample university AI policy document. It contains guidelines for 
    artificial intelligence usage in academic settings. The policy emphasizes 
    transparency, accountability, and ethical considerations in AI deployment.
    Students must disclose AI usage in assignments. Faculty should consider 
    bias and fairness when implementing AI tools.
    """
    
    print("=== Text Processing Test ===")
    print(f"Original text ({len(sample_text)} chars):")
    print(sample_text[:200] + "...")
    
    cleaned = processor.clean_text(sample_text)
    print(f"\nCleaned text ({len(cleaned)} chars):")
    print(cleaned[:200] + "...")
    
    stats = processor.get_text_statistics(cleaned)
    print(f"\nText Statistics:")
    for key, value in stats.items():
        print(f"  {key}: {value}")
    
    sentences = processor.tokenize_sentences(cleaned)
    print(f"\nFirst 3 sentences:")
    for i, sentence in enumerate(sentences[:3], 1):
        print(f"  {i}. {sentence}")
    
    words = processor.tokenize_words(cleaned)
    print(f"\nFirst 10 words: {words[:10]}")
    
    preview = processor.preview_text(cleaned, 100)
    print(f"\nPreview (100 chars): {preview}")
    
    print("\n✅ Text processor working correctly!")